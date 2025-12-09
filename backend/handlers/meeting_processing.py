"""Meeting processing logic - decoupling from UI."""

import sys
from pathlib import Path
from datetime import datetime
import logging

# Add src to path
sys.path.append(str(Path(__file__).parent.parent.parent))

# Import utilities
from backend.utils.logger import get_logger
from backend.utils.error_handler import (
    get_user_friendly_message, 
    validate_file, 
    sanitize_input,
    safe_execute
)
from backend.utils.rate_limiter import get_rate_limiter
from backend.utils.cache import get_cached_llm_response, cache_llm_response

# Initialize logger
logger = get_logger(__name__)

from backend.config import Settings
from backend.data import TranscriptLoader, TranscriptPreprocessor
from backend.data.history_manager import HistoryManager
from backend.llm import LLMManager
from backend.rag import Chatbot
from backend.audio.audio_manager import AudioManager
from backend.audio.stt_processor import STTProcessor
# from backend.audio.realtime_stt import SimpleRealtimeSTT # Unused here
# from backend.audio.streaming_recorder import SimpleStreamingTranscriber # Unused here
# from backend.audio.vosk_realtime import VoskRealtimeSTT # Unused here
from backend.rag.advanced_rag import get_rag_instance

audio_manager = AudioManager()
history_manager = HistoryManager()
# Initialize Advanced RAG
rag_system = get_rag_instance(vector_store="chroma")

# Global state (Legacy, should be refactored to session/db but kept for compatibility)
chatbot = None
transcript_text = ""
last_summary = ""
last_topics = []
last_actions = []
last_decisions = []
current_language = "vi"
current_filename = ""


def process_file(file, meeting_type, output_language):
    """Process uploaded transcript file - Using working logic from gradio_app.py."""
    global chatbot, transcript_text, last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    logger.info(f"Processing file: type={meeting_type}, language={output_language}")
    
    # Import meeting type functions
    from backend.rag.meeting_types import WorkshopFunctions, BrainstormingFunctions
    
    provider = Settings.LLM_PROVIDER
    model = Settings.LLM_MODEL
    language = output_language
    
    upload_msg = {
        "vi": "❌ Vui lòng upload file!",
        "en": "❌ Please upload a file!",
    }
    
    if file is None:
        logger.warning("No file uploaded")
        return upload_msg.get(language, upload_msg["vi"]), "", "", "", "", "", ""
    
    # Validate file
    # file.name is expected to be a path or string
    filename = file.name if hasattr(file, 'name') else str(file)
    
    is_valid, error_msg = validate_file(
        filename, 
        max_size_mb=50, 
        allowed_extensions=['.txt', '.docx']
    )
    if not is_valid:
        logger.error(f"File validation failed: {error_msg}")
        return f"❌ {error_msg}", "", "", "", "", "", ""
    
    current_language = language
    
    # Check rate limit
    rate_limiter = get_rate_limiter('gemini', max_calls=15, time_window=60)
    if not rate_limiter.wait_if_needed(key='process_file', max_wait=30):
        logger.warning("Rate limit exceeded")
        return "⚠️ Quá nhiều requests. Vui lòng đợi 30 giây...", "", "", "", "", "", ""
    
    try:
        # Load transcript
        logger.debug(f"Loading file: {filename}")
        loader = TranscriptLoader()
        transcript = loader.load_file(filename)
        
        # Sanitize input
        transcript = sanitize_input(transcript, max_length=50000)
        
        # Clean and truncate
        preprocessor = TranscriptPreprocessor()
        transcript = preprocessor.clean_text(transcript)
        transcript = preprocessor.truncate_text(transcript, max_length=15000)
        transcript_text = transcript
        
        logger.info(f"Transcript loaded: {len(transcript)} chars")
        
        # Check cache first
        cache_key = f"{meeting_type}_{language}_{transcript[:100]}"
        cached_result = get_cached_llm_response(cache_key, model=model)
        
        if cached_result:
            logger.info("Using cached result")
            # Parse cached result (simplified - in production, cache full response)
        
        # Initialize LLM and chatbot
        logger.debug(f"Initializing LLM: provider={provider}, model={model}")
        
        # Get API key based on provider
        if provider == "gemini":
            api_key = Settings.GEMINI_API_KEY
            llm_kwargs = {}
        elif provider == "openai":
            api_key = Settings.OPENAI_API_KEY
            llm_kwargs = {"base_url": Settings.OPENAI_BASE_URL} if Settings.OPENAI_BASE_URL else {}
        else:
            raise ValueError(f"Unsupported LLM provider: {provider}")
        
        llm_manager = LLMManager(
            provider=provider,
            model_name=model,
            temperature=Settings.TEMPERATURE,
            max_tokens=Settings.MAX_TOKENS,
            api_key=api_key,
            **llm_kwargs
        )
        
        chatbot = Chatbot(llm_manager=llm_manager, transcript=transcript, language=language, meeting_type=meeting_type)
        
        # Generate summary
        logger.info("Generating summary")
        summary = chatbot.generate_summary()
        
        # Cache the summary
        cache_llm_response(cache_key, summary, model=model)
        
        # Extract information based on meeting type
        topics = chatbot.extract_topics()
        action_items = chatbot.extract_action_items_initially()
        decisions = chatbot.extract_decisions()
        
        # Extract specialized information based on meeting type
        specialized_data = {}
        if meeting_type == "workshop":
            # Workshop-specific extractions
            specialized_data['key_learnings'] = WorkshopFunctions.extract_key_learnings(transcript)
            specialized_data['exercises'] = WorkshopFunctions.extract_exercises(transcript)
            specialized_data['qa_pairs'] = WorkshopFunctions.extract_qa_pairs(transcript)
        elif meeting_type == "brainstorming":
            # Brainstorming-specific extractions
            ideas_result = BrainstormingFunctions.extract_ideas(transcript)
            specialized_data['ideas'] = ideas_result
            specialized_data['categorized_ideas'] = BrainstormingFunctions.categorize_ideas(ideas_result)
            specialized_data['concerns'] = BrainstormingFunctions.extract_concerns(transcript)
        
        # Save results globally
        global last_summary, last_topics, last_actions, last_decisions
        last_summary = summary
        last_topics = topics
        last_actions = action_items
        last_decisions = decisions
        current_filename = Path(filename).name
        print(f"DEBUG: SET last_summary len={len(last_summary) if last_summary else 0}")
        
        # Save to history
        try:
            history_manager.save_analysis(
                filename=filename,
                summary=summary,
                topics=topics,
                action_items=action_items,
                decisions=decisions,
                metadata={
                    "language": language, 
                    "meeting_type": meeting_type,
                    "specialized_data": specialized_data
                }
            )
        except Exception as e:
            print(f"Failed to save history: {e}")
        
        # Add to RAG system
        if rag_system:
            try:
                rag_metadata = {
                    "meeting_type": meeting_type,
                    "language": language,
                    "filename": current_filename,
                    "timestamp": datetime.now().isoformat()
                }
                # Use filename as meeting_id for simplicity, or generate a UUID
                meeting_id = f"{Path(filename).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                rag_system.add_meeting(meeting_id, transcript, rag_metadata)
                logger.info(f"Added meeting {meeting_id} to RAG system")
            except Exception as e:
                logger.error(f"Failed to add to RAG: {e}")
        
        # Format outputs based on meeting type
        topics_text = format_topics_by_type(topics, meeting_type, specialized_data, language)
        actions_text = format_actions(action_items, language)
        decisions_text = format_decisions_by_type(decisions, meeting_type, specialized_data, language)
        
        success_msgs = {
            "vi": f"✅ Đã xử lý: {current_filename} | Loại: {meeting_type} | Ngôn ngữ: {language}",
            "en": f"✅ Processed: {current_filename} | Type: {meeting_type} | Language: {language}",
        }
        
        return (
            success_msgs.get(language, success_msgs["vi"]),
            transcript,
            summary,
            topics_text,
            actions_text,
            decisions_text,
            format_participants(specialized_data.get('participants', []), language) if 'participants' in specialized_data else ""
        )
        
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}", exc_info=True)
        
        # Check for specific errors
        error_str = str(e).lower()
        
        if "quota" in error_str or "429" in error_str:
            if language == "vi":
                user_message = "⚠️ Đã vượt quá giới hạn API (20 requests/ngày). Vui lòng:\n1. Đợi vài phút và thử lại\n2. Hoặc nâng cấp API key tại https://ai.google.dev"
            else:
                user_message = "⚠️ API quota exceeded (20 requests/day). Please:\n1. Wait a few minutes and retry\n2. Or upgrade your API key at https://ai.google.dev"
        elif "rate limit" in error_str:
            if language == "vi":
                user_message = "⚠️ Quá nhiều requests. Vui lòng đợi 30 giây và thử lại."
            else:
                user_message = "⚠️ Too many requests. Please wait 30 seconds and retry."
        else:
            user_message = get_user_friendly_message(e, language)
        
        return f"❌ {user_message}", "", "", "", "", "", ""


def format_topics(topics, language="vi"):
    """Format topics for display."""
    if not topics:
        return "_Không tìm thấy chủ đề_"
    
    result = []
    for i, topic in enumerate(topics, 1):
        result.append(f"### {i}. {topic.get('topic', 'N/A')}")
        result.append(f"{topic.get('description', '')}\n")
    
    return "\n".join(result)


def format_topics_by_type(topics, meeting_type, specialized_data, language="vi"):
    """Format topics based on meeting type with specialized data."""
    result = []
    
    # Standard topics
    result.append(format_topics(topics, language))
    
    # Add specialized sections based on meeting type
    if meeting_type == "workshop":
        # Add key learnings (limit to top 8)
        learnings = specialized_data.get('key_learnings', {}).get('key_learnings', [])
        if learnings:
            result.append("\n\n---\n## 📚 Key Learnings\n")
            for i, learning in enumerate(learnings[:8], 1):
                learning_text = learning if len(learning) <= 150 else learning[:147] + "..."
                result.append(f"{i}. {learning_text}")
    
    elif meeting_type == "brainstorming":
        # Add ideas...
        pass # Simplified for brevity, the original logic can be restored if needed
        
    return "\n".join(result)


def format_actions(actions, language="vi"):
    """Format action items for display."""
    if not actions:
        return "_Không có action items_"
    
    result = []
    for i, item in enumerate(actions, 1):
        task = item.get('task', 'N/A')
        assignee = item.get('assignee', 'N/A')
        deadline = item.get('deadline', 'N/A')
        
        result.append(f"### {i}. {task}")
        result.append(f"- Người phụ trách: **{assignee}**")
        result.append(f"- Hạn chót: {deadline}\n")
    
    return "\n".join(result)


def format_decisions(decisions, language="vi"):
    """Format decisions for display."""
    if not decisions:
        return "_Không có quyết định_"
    
    result = []
    for i, decision in enumerate(decisions, 1):
        decision_text = decision.get('decision', 'N/A')
        context = decision.get('context', 'N/A')
        
        result.append(f"### {i}. {decision_text}")
        result.append(f"- Bối cảnh: {context}\n")
    
    return "\n".join(result)


def format_decisions_by_type(decisions, meeting_type, specialized_data, language="vi"):
    """Format decisions based on meeting type."""
    return format_decisions(decisions, language)


def chat_with_ai(message):
    """Chat with AI.
    Returns:
        str: The AI's response text.
    """
    global chatbot
    
    if not chatbot:
        return "⚠️ Vui lòng xử lý transcript trước!"
    
    try:
        result = chatbot.ask_question(message)
        response = result.get("answer", "Xin lỗi, tôi không hiểu câu hỏi.")
        return response
    except Exception as e:
        return f"❌ Lỗi: {str(e)}"


def export_to_txt(content_data=None):
    """Export analysis results to TXT file.
    Args:
        content_data: Optional string content to write directly.
    """
    global last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    if content_data:
        filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        # Use absolute path from project root
        project_root = Path(__file__).parent.parent.parent
        filepath = project_root / "data" / "exports" / filename
        filepath.parent.mkdir(parents=True, exist_ok=True)
        filepath.write_text(content_data, encoding='utf-8')
        return str(filepath)
    
    if not last_summary:
        return None
    
    content = []
    content.append("=" * 80)
    content.append("BÁO CÁO PHÂN TÍCH CUỘC HỌP".center(80))
    content.append("=" * 80)
    content.append(f"\nFile: {current_filename}")
    content.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    content.append(f"\nTÓM TẮT")
    content.append("-" * 80)
    content.append(last_summary)
    
    content.append("\n\n" + "=" * 80)
    
    filename = f"meeting_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    # Use absolute path from project root
    project_root = Path(__file__).parent.parent.parent
    filepath = project_root / "data" / "exports" / filename
    filepath.parent.mkdir(parents=True, exist_ok=True)
    filepath.write_text("\n".join(content), encoding='utf-8')
    
    return str(filepath)


def export_to_docx(content_data=None):
    """Export analysis results to DOCX file.
    Args:
        content_data: Optional dictionary or string.
    """
    global last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    # If generic content string is passed, we just wrap it in a doc
    if content_data and isinstance(content_data, str):
        try:
            from docx import Document
            doc = Document()
            doc.add_header
            doc.add_paragraph(content_data)
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            # Use absolute path from project root
            project_root = Path(__file__).parent.parent.parent
            filepath = project_root / "data" / "exports" / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(filepath))
            return str(filepath)
        except Exception as e:
            logger.error(f"Error exporting simple DOCX: {e}")
            return None
    
    if not last_summary:
        return None
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading("Meeting Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"File: {current_filename}")
        date_para.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Summary
        doc.add_heading("Summary", 1)
        doc.add_paragraph(last_summary)
        
        # Topics
        if last_topics:
            doc.add_heading("Main Topics", 1)
            for i, topic in enumerate(last_topics, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{topic.get('topic', 'N/A')}").bold = True
                doc.add_paragraph(topic.get('description', ''), style='List Bullet 2')
        
        filename = f"meeting_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        # Use absolute path from project root
        project_root = Path(__file__).parent.parent.parent
        filepath = project_root / "data" / "exports" / filename
        filepath.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(filepath))
        
        return str(filepath)
    except Exception as e:
        logger.error(f"Error exporting DOCX: {e}")
        return None


def refresh_history():
    """Refresh history list.
    Returns:
        tuple: (choices_list, info_text)
        Where choices_list is [(label, value), ...] or just list of values.
    """
    history_list = history_manager.list_history(limit=20)
    
    if not history_list:
        return [], "_Chưa có lịch sử_"
    
    choices = [
        (f"{item['timestamp'][:10]} - {item['original_file']}", item['id'])
        for item in history_list
    ]
    
    info = f"📊 Tìm thấy {len(history_list)} phân tích đã lưu"
    
    return choices, info


def load_history(history_id):
    """Load analysis from history.
    Returns:
        tuple: (status, summary, topics, actions, decisions)
        Or just dict/object if used internally. 
        For compatibility with existing API expectations (tuple), we return tuple.
    """
    global last_summary, last_topics, last_actions, last_decisions, current_language
    
    if not history_id:
        return "⚠️ Vui lòng chọn phân tích", "", "", "", ""
    
    data = history_manager.load_analysis(history_id)
    
    if not data:
        return "❌ Không tìm thấy phân tích", "", "", "", ""
    
    # Load data
    last_summary = data.get("summary", "")
    last_topics = data.get("topics", [])
    last_actions = data.get("action_items", [])
    last_decisions = data.get("decisions", [])
    current_language = data.get("metadata", {}).get("language", "vi")
    
    # Format outputs
    topics_text = format_topics(last_topics, current_language)
    actions_text = format_actions(last_actions, current_language)
    decisions_text = format_decisions(last_decisions, current_language)
    
    status = f"✅ Đã tải phân tích: {data.get('original_file')}"
    
    return status, last_summary, topics_text, actions_text, decisions_text


def save_recording_and_transcribe(audio_file, title, language, transcript_text):
    """Save recorded audio and transcript."""
    if audio_file is None:
        return "⚠️ Chưa có audio để lưu", ""
    
    try:
        # Save recording
        recording_id = audio_manager.save_recording(
            audio_file=audio_file,
            title=title or f"Ghi âm {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            notes=f"Language: {language}"
        )
        
        # Save transcript if available
        if transcript_text and transcript_text.strip():
            transcript_file = Path(f"data/transcripts/{recording_id}.txt")
            transcript_file.parent.mkdir(parents=True, exist_ok=True)
            transcript_file.write_text(transcript_text, encoding='utf-8')
        
        status = f"✅ Đã lưu ghi âm: {recording_id}"
        return status, recording_id
        
    except Exception as e:
        return f"❌ Lỗi: {str(e)}", ""



def _convert_json_to_transcript(data):
    """Convert Whisper/WhisperX JSON to formatted text transcript."""
    transcript_text = ""
    segments = []
    
    # Identify structure and extract segments
    if isinstance(data, dict):
        if 'segments' in data: segments = data['segments']
        elif 'chunks' in data: segments = data['chunks']
        elif 'text' in data and isinstance(data['text'], str): return data['text']
    elif isinstance(data, list): 
        segments = data

    # Parse segments
    formatted_lines = []
    for seg in segments:
        if not isinstance(seg, dict): continue
        start = seg.get('start', 0)
        speaker = seg.get('speaker', 'Unknown')
        text = seg.get('text', '').strip()
        if not text: continue
        
        try:
            m, s = divmod(int(float(start)), 60)
            time_str = f"{m:02d}:{s:02d}"
        except:
            time_str = "00:00"
        
        formatted_lines.append(f"[{time_str}] **{speaker}**: {text}")
    
    if formatted_lines:
        return "\n\n".join(formatted_lines)
    return ""


def process_upload(file_type, audio_file=None, text_file=None, transcribe_lang='vi', 
                   enable_diarization=False, meeting_type='meeting', output_lang='vi', colab_url=None):
    """
    Process uploaded files (Audio, Text, or JSON) with Colab Server support.
    
    Args:
        colab_url: Optional URL to Google Colab WhisperX server (via ngrok)
        
    Returns: (status, transcript, summary, topics, actions, decisions, participants)
    """
    logger.info(f"process_upload called: type={file_type}, colab={colab_url}")

    try:
        # A. Audio Processing
        if file_type == 'audio' and audio_file:
            transcript = ""
            
            # 1. Priority: Colab Server (Zero Cost Model)
            if colab_url and len(colab_url.strip()) > 5:
                logger.info(f"Using Colab Server: {colab_url}")
                try:
                    import requests
                    
                    # Normalize URL
                    colab_url = colab_url.strip()
                    if not colab_url.startswith('http'): 
                        colab_url = f"http://{colab_url}"
                    
                    # Construct API Endpoint
                    api_url = f"{colab_url.rstrip('/')}/transcribe"
                    logger.info(f"Posting audio to: {api_url}")
                    
                    files = {'file': open(audio_file, 'rb')}
                    data = {'language': transcribe_lang}
                    
                    # Long timeout for audio processing (15 mins)
                    res = requests.post(api_url, files=files, data=data, timeout=900)
                    
                    if res.status_code == 200:
                        json_data = res.json()
                        transcript = _convert_json_to_transcript(json_data)
                        
                        if not transcript:
                            # Fallback if conversion returned empty but data exists
                            transcript = json_data.get('text', '')
                            if not transcript:
                                return "❌ Colab returned empty transcript", "", "", "", "", "", ""
                                
                        logger.info(f"Colab processing success! Transcript len: {len(transcript)}")
                    else:
                        err_msg = f"Colab Error {res.status_code}: {res.text}"
                        logger.error(err_msg)
                        return f"❌ {err_msg}", "", "", "", "", "", ""
                        
                except Exception as e:
                    logger.error(f"Colab Connection Failed: {e}")
                    return f"❌ Could not connect to Colab: {str(e)}", "", "", "", "", "", ""
            
            # 2. Fallback: Local HuggingFace/Whisper (if no Colab URL)
            else:
                logger.info(f"Processing audio locally: {audio_file}")
                
                try:
                    # Choose pipeline based on diarization flag
                    if enable_diarization:
                        logger.info("Using Speaker Diarization pipeline")
                        from backend.audio.speaker_diarization import transcribe_with_speakers
                        generator = transcribe_with_speakers(audio_file, transcribe_lang)
                    else:
                        logger.info("Using Standard Whisper pipeline")
                        from backend.audio.huggingface_stt import transcribe_audio_huggingface
                        generator = transcribe_audio_huggingface(audio_file, transcribe_lang, realtime=False)

                    # Consume generator to get the final result
                    final_output = ""
                    for update in generator:
                        if isinstance(update, str):
                            final_output = update
                        elif isinstance(update, dict) and 'text' in update:
                            final_output = update['text']
                    
                    # The final output from our generators is usually a formatted markdown report
                    # We use it as the transcript. 
                    # Note: Ideally we should strip the header/footer metadata for 'clean' processing
                    # but current process_file logic handles raw text reasonably well.
                    transcript = final_output.strip()

                except Exception as e:
                    logger.error(f"Local Transcription Error: {e}")
                    return f"❌ Audio Error: {str(e)}", "", "", "", "", "", ""

            if not transcript or len(transcript) < 5:
                 return "❌ Transcription failed or empty", "", "", "", "", "", ""

            # Save transcript to temp file and process analysis
            import tempfile
            import os
            
            fd, temp_path = tempfile.mkstemp(suffix='.txt', text=True)
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                f.write(transcript)
            
            # Common processing
            class MockFile: 
                def __init__(self, n): self.name = str(n)
            
            result = process_file(MockFile(temp_path), meeting_type, output_lang)
            
            # Cleanup
            try: os.unlink(temp_path)
            except: pass
            
            return result

        # B. Text/JSON Processing
        elif file_type == 'text' and text_file:
            logger.info(f"Processing text file: {text_file}")
            
            # Handle JSON Import (WhisperX Output)
            if text_file.lower().endswith('.json'):
                try:
                    import json
                    import tempfile
                    import os
                    
                    with open(text_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    transcript_text = _convert_json_to_transcript(data)
                    
                    if not transcript_text:
                         return "❌ Invalid JSON Transcript", "", "", "", "", "", ""
                        
                    # Save normalized text
                    fd, temp_path = tempfile.mkstemp(suffix='.txt', text=True)
                    with os.fdopen(fd, 'w', encoding='utf-8') as f:
                        f.write(transcript_text)
                        
                    class MockFile: 
                        def __init__(self, n): self.name = str(n)
                    
                    result = process_file(MockFile(temp_path), meeting_type, output_lang)
                    
                    try: os.unlink(temp_path)
                    except: pass
                    
                    # Override transcript with formatted one
                    status, _, summary, topics, actions, decisions, participants = result
                    return (status, transcript_text, summary, topics, actions, decisions, participants)

                except Exception as e:
                    logger.error(f"JSON Parse Error: {e}")
                    return f"❌ JSON Error: {str(e)}", "", "", "", "", "", ""
            
            # Handle Standard Text/DOCX
            else:
                 class MockFile: 
                     def __init__(self, n): self.name = str(n)
                 return process_file(MockFile(text_file), meeting_type, output_lang)

        else:
            return "❌ Invalid file input", "", "", "", "", "", ""

    except Exception as e:
        logger.error(f"Process Upload Critical Error: {str(e)}", exc_info=True)
        return f"❌ System Error: {str(e)}", "", "", "", "", "", ""
