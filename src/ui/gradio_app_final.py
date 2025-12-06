"""Final Gradio UI - Best of both worlds: Working logic + Beautiful design."""

import gradio as gr
import sys
from pathlib import Path
from datetime import datetime

# Add src to path
sys.path.append(str(Path(__file__).parent.parent.parent))

# Import utilities
from src.utils.logger import get_logger
from src.utils.error_handler import (
    get_user_friendly_message, 
    validate_file, 
    sanitize_input,
    safe_execute
)
from src.utils.rate_limiter import get_rate_limiter
from src.utils.cache import get_cached_llm_response, cache_llm_response

# Initialize logger
logger = get_logger(__name__)

from src.config import Settings
from src.data import TranscriptLoader, TranscriptPreprocessor
from src.data.history_manager import HistoryManager
from src.llm import LLMManager
from src.rag import Chatbot
from src.audio.audio_manager import AudioManager
from src.audio.stt_processor import STTProcessor
from src.audio.realtime_stt import SimpleRealtimeSTT
from src.audio.streaming_recorder import SimpleStreamingTranscriber
from src.audio.vosk_realtime import VoskRealtimeSTT
from src.rag.advanced_rag import get_rag_instance

audio_manager = AudioManager()
history_manager = HistoryManager()
# Initialize Advanced RAG
rag_system = get_rag_instance(vector_store="chroma")
stt_processor = None  # Lazy init (requires API key)
realtime_stt = SimpleRealtimeSTT()
streaming_transcriber = SimpleStreamingTranscriber(chunk_interval=10)  # Update every 10s
vosk_stt = VoskRealtimeSTT()  # Free, offline, realtime


def process_file(file, meeting_type, output_language):
    """Process uploaded transcript file - Using working logic from gradio_app.py."""
    global chatbot, transcript_text, last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    logger.info(f"Processing file: type={meeting_type}, language={output_language}")
    
    # Import meeting type functions
    from src.rag.meeting_types import WorkshopFunctions, BrainstormingFunctions
    
    provider = "gemini"
    model = "gemini-2.5-flash"
    language = output_language
    
    upload_msg = {
        "vi": "❌ Vui lòng upload file!",
        "en": "❌ Please upload a file!",
        "ja": "❌ ファイルをアップロードしてください！",
        "ko": "❌ 파일을 업로드해주세요!",
        "zh-CN": "❌ 请上传文件！",
        "es": "❌ Por favor sube un archivo!",
        "fr": "❌ Veuillez télécharger un fichier!",
        "de": "❌ Bitte laden Sie eine Datei hoch!"
    }
    
    if file is None:
        logger.warning("No file uploaded")
        return upload_msg.get(language, upload_msg["vi"]), "", "", "", ""
    
    # Validate file
    is_valid, error_msg = validate_file(
        file.name, 
        max_size_mb=50, 
        allowed_extensions=['.txt', '.docx']
    )
    if not is_valid:
        logger.error(f"File validation failed: {error_msg}")
        return f"❌ {error_msg}", "", "", "", ""
    
    current_language = language
    
    # Check rate limit
    rate_limiter = get_rate_limiter('gemini', max_calls=15, time_window=60)
    if not rate_limiter.wait_if_needed(key='process_file', max_wait=30):
        logger.warning("Rate limit exceeded")
        return "⚠️ Quá nhiều requests. Vui lòng đợi 30 giây...", "", "", "", ""
    
    try:
        # Load transcript
        logger.debug(f"Loading file: {file.name}")
        loader = TranscriptLoader()
        transcript = loader.load_file(file.name)
        
        # Sanitize input
        transcript = sanitize_input(transcript, max_length=50000)
        
        # Clean and truncate
        preprocessor = TranscriptPreprocessor()
        transcript = preprocessor.clean_text(transcript)
        transcript = preprocessor.truncate_text(transcript, max_length=15000)
        transcript_text = transcript
        
        logger.info(f"Transcript loaded: {len(transcript)} chars")
        
        # Check cache first
        cache_key = f"{meeting_type}_{language}_{transcript[:100]}"
        cached_result = get_cached_llm_response(cache_key, model=model)
        
        if cached_result:
            logger.info("Using cached result")
            # Parse cached result (simplified - in production, cache full response)
        
        # Initialize LLM and chatbot
        logger.debug("Initializing LLM")
        llm_manager = LLMManager(
            provider=provider,
            model_name=model,
            temperature=Settings.TEMPERATURE,
            max_tokens=Settings.MAX_TOKENS,
            api_key=Settings.GEMINI_API_KEY,
        )
        
        chatbot = Chatbot(llm_manager=llm_manager, transcript=transcript, language=language, meeting_type=meeting_type)
        
        # Generate summary
        logger.info("Generating summary")
        summary = chatbot.generate_summary()
        
        # Cache the summary
        cache_llm_response(cache_key, summary, model=model)
        
        # Extract information based on meeting type
        topics = chatbot.extract_topics()
        action_items = chatbot.extract_action_items_initially()
        decisions = chatbot.extract_decisions()
        
        # Extract specialized information based on meeting type
        specialized_data = {}
        if meeting_type == "workshop":
            # Workshop-specific extractions
            specialized_data['key_learnings'] = WorkshopFunctions.extract_key_learnings(transcript)
            specialized_data['exercises'] = WorkshopFunctions.extract_exercises(transcript)
            specialized_data['qa_pairs'] = WorkshopFunctions.extract_qa_pairs(transcript)
        elif meeting_type == "brainstorming":
            # Brainstorming-specific extractions
            ideas_result = BrainstormingFunctions.extract_ideas(transcript)
            specialized_data['ideas'] = ideas_result
            specialized_data['categorized_ideas'] = BrainstormingFunctions.categorize_ideas(ideas_result)
            specialized_data['concerns'] = BrainstormingFunctions.extract_concerns(transcript)
        
        # Save results globally
        last_summary = summary
        last_topics = topics
        last_actions = action_items
        last_decisions = decisions
        current_filename = Path(file.name).name
        
        # Save to history
        try:
            history_manager.save_analysis(
                filename=file.name,
                summary=summary,
                topics=topics,
                action_items=action_items,
                decisions=decisions,
                metadata={
                    "language": language, 
                    "meeting_type": meeting_type,
                    "specialized_data": specialized_data
                }
            )
        except Exception as e:
            print(f"Failed to save history: {e}")
        
        # Add to RAG system
        if rag_system:
            try:
                rag_metadata = {
                    "meeting_type": meeting_type,
                    "language": language,
                    "filename": current_filename,
                    "timestamp": datetime.now().isoformat()
                }
                # Use filename as meeting_id for simplicity, or generate a UUID
                meeting_id = f"{Path(file.name).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                rag_system.add_meeting(meeting_id, transcript, rag_metadata)
                logger.info(f"Added meeting {meeting_id} to RAG system")
            except Exception as e:
                logger.error(f"Failed to add to RAG: {e}")
        
        # Format outputs based on meeting type
        topics_text = format_topics_by_type(topics, meeting_type, specialized_data, language)
        actions_text = format_actions(action_items, language)
        decisions_text = format_decisions_by_type(decisions, meeting_type, specialized_data, language)
        
        success_msgs = {
            "vi": f"✅ Đã xử lý: {current_filename} | Loại: {meeting_type} | Ngôn ngữ: {language}",
            "en": f"✅ Processed: {current_filename} | Type: {meeting_type} | Language: {language}",
            "ja": f"✅ 処理完了: {current_filename} | タイプ: {meeting_type} | 言語: {language}",
            "ko": f"✅ 처리 완료: {current_filename} | 유형: {meeting_type} | 언어: {language}",
            "zh-CN": f"✅ 已处理: {current_filename} | 类型: {meeting_type} | 语言: {language}",
            "es": f"✅ Procesado: {current_filename} | Tipo: {meeting_type} | Idioma: {language}",
            "fr": f"✅ Traité: {current_filename} | Type: {meeting_type} | Langue: {language}",
            "de": f"✅ Verarbeitet: {current_filename} | Typ: {meeting_type} | Sprache: {language}"
        }
        
        return (
            success_msgs.get(language, success_msgs["vi"]),
            summary,
            topics_text,
            actions_text,
            decisions_text
        )
        
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}", exc_info=True)
        
        # Get user-friendly error message
        user_message = get_user_friendly_message(e, language)
        
        error_prefixes = {
            "vi": "❌ Lỗi:",
            "en": "❌ Error:",
            "ja": "❌ エラー:",
            "ko": "❌ 오류:",
            "zh-CN": "❌ 错误:",
            "es": "❌ Error:",
            "fr": "❌ Erreur:",
            "de": "❌ Fehler:"
        }
        
        return f"{error_prefixes.get(language, error_prefixes['vi'])} {user_message}", "", "", "", ""


def format_topics(topics, language="vi"):
    """Format topics for display."""
    labels = {
        "vi": {"no_data": "_Không tìm thấy chủ đề_", "description": "Mô tả"},
        "en": {"no_data": "_No topics found_", "description": "Description"},
        "ja": {"no_data": "_トピックが見つかりません_", "description": "説明"},
        "ko": {"no_data": "_주제를 찾을 수 없습니다_", "description": "설명"},
        "zh-CN": {"no_data": "_未找到主题_", "description": "描述"},
        "es": {"no_data": "_No se encontraron temas_", "description": "Descripción"},
        "fr": {"no_data": "_Aucun sujet trouvé_", "description": "Description"},
        "de": {"no_data": "_Keine Themen gefunden_", "description": "Beschreibung"}
    }
    lang = labels.get(language, labels["vi"])
    
    if not topics:
        return lang["no_data"]
    
    result = []
    for i, topic in enumerate(topics, 1):
        result.append(f"### {i}. {topic.get('topic', 'N/A')}")
        result.append(f"{topic.get('description', '')}\n")
    
    return "\n".join(result)


def format_topics_by_type(topics, meeting_type, specialized_data, language="vi"):
    """Format topics based on meeting type with specialized data."""
    result = []
    
    # Standard topics
    result.append(format_topics(topics, language))
    
    # Add specialized sections based on meeting type
    if meeting_type == "workshop":
        # Add key learnings (limit to top 8)
        learnings = specialized_data.get('key_learnings', {}).get('key_learnings', [])
        if learnings:
            result.append("\n\n---\n## 📚 Key Learnings\n")
            for i, learning in enumerate(learnings[:8], 1):
                # Truncate long learnings
                learning_text = learning if len(learning) <= 150 else learning[:147] + "..."
                result.append(f"{i}. {learning_text}")
            
            if len(learnings) > 8:
                result.append(f"\n_...and {len(learnings) - 8} more learnings_")
            
            result.append(f"\n**Total: {len(learnings)} key learnings**")
        
        # Add exercises (limit to top 10)
        exercises = specialized_data.get('exercises', {}).get('exercises', [])
        if exercises:
            result.append("\n\n---\n## 🎯 Exercises & Activities\n")
            for i, ex in enumerate(exercises[:10], 1):
                ex_title = ex.get('title', 'N/A')
                # Truncate long titles
                if len(ex_title) > 100:
                    ex_title = ex_title[:97] + "..."
                result.append(f"{i}. {ex_title}")
            
            if len(exercises) > 10:
                result.append(f"\n_...and {len(exercises) - 10} more exercises_")
            
            result.append(f"\n**Total: {len(exercises)} exercises**")
        
        # Add Q&A (limit to top 5 pairs)
        qa_pairs = specialized_data.get('qa_pairs', {}).get('qa_pairs', [])
        if qa_pairs:
            result.append("\n\n---\n## ❓ Q&A Highlights\n")
            for i, qa in enumerate(qa_pairs[:5], 1):
                question = qa.get('question', 'N/A')
                answer = qa.get('answer', 'N/A')
                
                # Truncate long Q&A
                if len(question) > 100:
                    question = question[:97] + "..."
                if len(answer) > 150:
                    answer = answer[:147] + "..."
                
                result.append(f"**Q{i}:** {question}")
                result.append(f"**A{i}:** {answer}\n")
            
            if len(qa_pairs) > 5:
                result.append(f"_...and {len(qa_pairs) - 5} more Q&A pairs_")
            
            result.append(f"\n**Total: {len(qa_pairs)} Q&A pairs**")
    
    elif meeting_type == "brainstorming":
        # Add ideas by category (limit to top ideas per category)
        categorized = specialized_data.get('categorized_ideas', {}).get('categorized_ideas', {})
        if categorized:
            result.append("\n\n---\n## 💡 Top Ideas by Category\n")
            
            # Count total ideas
            total_ideas = sum(len(ideas) for ideas in categorized.values())
            
            for category, ideas in categorized.items():
                if ideas and len(ideas) > 0:
                    # Show top 5 ideas per category
                    top_ideas = ideas[:5]
                    result.append(f"\n### {category} ({len(ideas)} ideas)")
                    for i, idea in enumerate(top_ideas, 1):
                        idea_text = idea.get('idea', 'N/A')
                        # Truncate long ideas
                        if len(idea_text) > 100:
                            idea_text = idea_text[:97] + "..."
                        result.append(f"{i}. {idea_text}")
                    
                    # Show "and X more" if there are more ideas
                    if len(ideas) > 5:
                        result.append(f"   _...and {len(ideas) - 5} more ideas_")
            
            result.append(f"\n**Total: {total_ideas} ideas generated**")
        
        # Add concerns (limit to top 5)
        concerns = specialized_data.get('concerns', {}).get('concerns', [])
        if concerns:
            result.append("\n\n---\n## ⚠️ Key Concerns & Challenges\n")
            for i, concern in enumerate(concerns[:5], 1):
                # Truncate long concerns
                concern_text = concern if len(concern) <= 150 else concern[:147] + "..."
                result.append(f"{i}. {concern_text}")
            
            if len(concerns) > 5:
                result.append(f"\n_...and {len(concerns) - 5} more concerns_")
    
    return "\n".join(result)


def format_actions(actions, language="vi"):
    """Format action items for display."""
    labels = {
        "vi": {
            "no_data": "_Không có action items_", 
            "assignee": "👤 Người phụ trách", 
            "deadline": "📅 Hạn chót",
            "not_assigned": "Chưa phân công",
            "not_specified": "Chưa xác định"
        },
        "en": {
            "no_data": "_No action items_", 
            "assignee": "👤 Assignee", 
            "deadline": "📅 Deadline",
            "not_assigned": "Not assigned",
            "not_specified": "Not specified"
        },
        "ja": {
            "no_data": "_アクションアイテムなし_", 
            "assignee": "👤 担当者", 
            "deadline": "📅 期限",
            "not_assigned": "未割り当て",
            "not_specified": "未指定"
        },
        "ko": {
            "no_data": "_액션 아이템 없음_", 
            "assignee": "👤 담당자", 
            "deadline": "📅 마감일",
            "not_assigned": "미할당",
            "not_specified": "미지정"
        },
        "zh-CN": {
            "no_data": "_无行动项_", 
            "assignee": "👤 负责人", 
            "deadline": "📅 截止日期",
            "not_assigned": "未分配",
            "not_specified": "未指定"
        },
        "es": {
            "no_data": "_Sin elementos de acción_", 
            "assignee": "👤 Responsable", 
            "deadline": "📅 Fecha límite",
            "not_assigned": "No asignado",
            "not_specified": "No especificado"
        },
        "fr": {
            "no_data": "_Aucun élément d'action_", 
            "assignee": "👤 Responsable", 
            "deadline": "📅 Date limite",
            "not_assigned": "Non assigné",
            "not_specified": "Non spécifié"
        },
        "de": {
            "no_data": "_Keine Aktionselemente_", 
            "assignee": "👤 Verantwortlich", 
            "deadline": "📅 Frist",
            "not_assigned": "Nicht zugewiesen",
            "not_specified": "Nicht angegeben"
        }
    }
    lang = labels.get(language, labels["vi"])
    
    if not actions:
        return lang["no_data"]
    
    result = []
    for i, item in enumerate(actions, 1):
        task = item.get('task', 'N/A')
        assignee = item.get('assignee', lang['not_assigned'])
        deadline = item.get('deadline', lang['not_specified'])
        
        result.append(f"### {i}. {task}")
        result.append(f"- {lang['assignee']}: **{assignee}**")
        result.append(f"- {lang['deadline']}: {deadline}\n")
    
    return "\n".join(result)


def format_decisions(decisions, language="vi"):
    """Format decisions for display."""
    labels = {
        "vi": {
            "no_data": "_Không có quyết định_", 
            "context": "📝 Bối cảnh",
            "no_context": "Không có bối cảnh"
        },
        "en": {
            "no_data": "_No decisions_", 
            "context": "📝 Context",
            "no_context": "No context"
        },
        "ja": {
            "no_data": "_決定なし_", 
            "context": "📝 文脈",
            "no_context": "文脈なし"
        },
        "ko": {
            "no_data": "_결정 없음_", 
            "context": "📝 맥락",
            "no_context": "맥락 없음"
        },
        "zh-CN": {
            "no_data": "_无决定_", 
            "context": "📝 背景",
            "no_context": "无背景"
        },
        "es": {
            "no_data": "_Sin decisiones_", 
            "context": "📝 Contexto",
            "no_context": "Sin contexto"
        },
        "fr": {
            "no_data": "_Aucune décision_", 
            "context": "📝 Contexte",
            "no_context": "Pas de contexte"
        },
        "de": {
            "no_data": "_Keine Entscheidungen_", 
            "context": "📝 Kontext",
            "no_context": "Kein Kontext"
        }
    }
    lang = labels.get(language, labels["vi"])
    
    if not decisions:
        return lang["no_data"]
    
    result = []
    for i, decision in enumerate(decisions, 1):
        decision_text = decision.get('decision', 'N/A')
        context = decision.get('context', lang['no_context'])
        
        result.append(f"### {i}. {decision_text}")
        result.append(f"- {lang['context']}: {context}\n")
    
    return "\n".join(result)


def format_decisions_by_type(decisions, meeting_type, specialized_data, language="vi"):
    """Format decisions based on meeting type."""
    # Just return standard decisions, no extra sections
    return format_decisions(decisions, language)


def chat_with_ai(message, history):
    """Chat with AI."""
    global chatbot
    
    if history is None:
        history = []
    
    if not chatbot:
        new_history = history.copy()
        new_history.append({"role": "user", "content": message})
        new_history.append({"role": "assistant", "content": "⚠️ Vui lòng xử lý transcript trước!"})
        return new_history
    
    try:
        result = chatbot.ask_question(message)
        response = result.get("answer", "Xin lỗi, tôi không hiểu câu hỏi.")
        
        # Use messages format (not tuples)
        new_history = history.copy() if history else []
        new_history.append({"role": "user", "content": message})
        new_history.append({"role": "assistant", "content": response})
        return new_history
    except Exception as e:
        new_history = history.copy() if history else []
        new_history.append({"role": "user", "content": message})
        new_history.append({"role": "assistant", "content": f"❌ Lỗi: {str(e)}"})
        return new_history


def export_to_txt():
    """Export analysis results to TXT file."""
    global last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    if not last_summary:
        return None
    
    labels = {
        "vi": {"title": "BÁO CÁO PHÂN TÍCH CUỘC HỌP", "summary": "TÓM TẮT", "topics": "CHỦ ĐỀ", "actions": "ACTION ITEMS", "decisions": "QUYẾT ĐỊNH"},
        "en": {"title": "MEETING ANALYSIS REPORT", "summary": "SUMMARY", "topics": "TOPICS", "actions": "ACTION ITEMS", "decisions": "DECISIONS"}
    }
    lang = labels.get(current_language, labels["vi"])
    
    content = []
    content.append("=" * 80)
    content.append(lang["title"].center(80))
    content.append("=" * 80)
    content.append(f"\nFile: {current_filename}")
    content.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    content.append("=" * 80)
    
    content.append(f"\n{lang['summary']}")
    content.append("-" * 80)
    content.append(last_summary)
    
    if last_topics:
        content.append(f"\n\n{lang['topics']}")
        content.append("-" * 80)
        for i, topic in enumerate(last_topics, 1):
            content.append(f"\n{i}. {topic.get('topic', 'N/A')}")
            content.append(f"   {topic.get('description', '')}")
    
    if last_actions:
        content.append(f"\n\n{lang['actions']}")
        content.append("-" * 80)
        for i, action in enumerate(last_actions, 1):
            content.append(f"\n{i}. {action.get('task', 'N/A')}")
            content.append(f"   Assignee: {action.get('assignee', 'N/A')}")
            content.append(f"   Deadline: {action.get('deadline', 'N/A')}")
    
    if last_decisions:
        content.append(f"\n\n{lang['decisions']}")
        content.append("-" * 80)
        for i, decision in enumerate(last_decisions, 1):
            content.append(f"\n{i}. {decision.get('decision', 'N/A')}")
            content.append(f"   Context: {decision.get('context', 'N/A')}")
    
    content.append("\n\n" + "=" * 80)
    
    filename = f"meeting_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    filepath = Path(filename)
    filepath.write_text("\n".join(content), encoding='utf-8')
    
    return str(filepath)


def refresh_history():
    """Refresh history dropdown."""
    history_list = history_manager.list_history(limit=20)
    
    if not history_list:
        return gr.Dropdown(choices=[], value=None), "_Chưa có lịch sử_"
    
    choices = [
        (f"{item['timestamp'][:10]} - {item['original_file']}", item['id'])
        for item in history_list
    ]
    
    info = f"📊 Tìm thấy {len(history_list)} phân tích đã lưu"
    
    return gr.Dropdown(choices=choices, value=None), info


def load_history(history_id):
    """Load analysis from history."""
    global last_summary, last_topics, last_actions, last_decisions, current_language
    
    if not history_id:
        return "⚠️ Vui lòng chọn phân tích từ danh sách", "", "", "", ""
    
    data = history_manager.load_analysis(history_id)
    
    if not data:
        return "❌ Không tìm thấy phân tích", "", "", "", ""
    
    # Load data
    last_summary = data.get("summary", "")
    last_topics = data.get("topics", [])
    last_actions = data.get("action_items", [])
    last_decisions = data.get("decisions", [])
    current_language = data.get("metadata", {}).get("language", "vi")
    
    # Format outputs
    topics_text = format_topics(last_topics, current_language)
    actions_text = format_actions(last_actions, current_language)
    decisions_text = format_decisions(last_decisions, current_language)
    
    status = f"✅ Đã tải phân tích: {data.get('original_file')} ({data.get('timestamp')[:10]})"
    
    return status, last_summary, topics_text, actions_text, decisions_text


def toggle_recording_modal():
    """Toggle recording modal visibility."""
    return gr.Group(visible=True)


def cancel_recording():
    """Cancel recording and hide modal."""
    return gr.Group(visible=False), "_Đã hủy_", gr.Audio(visible=False)


def start_recording_session(title, language, auto_translate, translate_to):
    """Start recording session."""
    status = f"""
### 🎙️ Đang ghi âm...

**Tiêu đề:** {title or f"Ghi âm {datetime.now().strftime('%d/%m/%Y')}"}  
**Ngôn ngữ:** {language}  
**Tự động dịch:** {'Có' if auto_translate else 'Không'}  
{f"**Dịch sang:** {translate_to}" if auto_translate else ""}

⏺️ Đang ghi... Nhấn Stop khi hoàn thành.
    """
    return status, gr.Audio(visible=True)


def save_recording_and_transcribe(audio_file, title, language, transcript_text):
    """Save recorded audio and transcript from main recording tab."""
    if audio_file is None:
        return "⚠️ Chưa có audio để lưu", ""
    
    try:
        # Save recording
        recording_id = audio_manager.save_recording(
            audio_file=audio_file,
            title=title or f"Ghi âm {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            notes=f"Language: {language}"
        )
        
        # Save transcript if available
        if transcript_text and transcript_text.strip():
            transcript_file = Path(f"data/transcripts/{recording_id}.txt")
            transcript_file.parent.mkdir(parents=True, exist_ok=True)
            transcript_file.write_text(transcript_text, encoding='utf-8')
        
        status = f"""✅ Đã lưu ghi âm và transcript!

**ID:** {recording_id}  
**Audio:** data/recordings/{recording_id}.wav
**Transcript:** data/transcripts/{recording_id}.txt

💡 Xem trong tab "Thư Viện Lưu Trữ > Lịch Sử Ghi Âm"
        """
        
        return status, recording_id
        
    except Exception as e:
        return f"❌ Lỗi: {str(e)}", ""


def transcribe_audio_whisper(audio_file, audio_language):
    """Transcribe audio using Whisper (local offline model).
    
    This is a generator function that yields progressive updates.
    """
    if audio_file is None or audio_file == "":
        yield "🎙️ Sẵn sàng ghi âm. Nhấn microphone icon để bắt đầu..."
        return
    
    try:
        import whisper
        import torch
        
        yield "🔄 Đang khởi tạo Whisper (Local Offline)..."
        
        # Check if CUDA is available
        device = "cuda" if torch.cuda.is_available() else "cpu"
        yield f"⚙️  Sử dụng: {device.upper()}"
        
        # Load model (base model - good balance between speed and accuracy)
        yield "📥 Đang tải Whisper model (base)..."
        model = whisper.load_model("base", device=device)
        
        yield f"🎤 Đang transcribe audio (ngôn ngữ: {audio_language})..."
        
        # Transcribe
        result = model.transcribe(
            audio_file,
            language=audio_language,
            fp16=(device == "cuda")  # Use FP16 only on GPU
        )
        
        final_transcript = result["text"].strip()
        
        # Get current date time
        from datetime import datetime
        now = datetime.now().strftime("%d/%m/%Y %H:%M")
        
        yield f"""📝 **Transcript ({audio_language}):** {now}

{final_transcript}
        """
            
    except ImportError:
        yield """❌ Lỗi: Chưa cài đặt Whisper

💡 **Cài đặt Whisper (Local Offline):**

```bash
pip install openai-whisper
```

**Hoặc với GPU support:**
```bash
pip install openai-whisper torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118
```

**Ưu điểm:**
- ✅ Miễn phí, không cần API key
- ✅ Hoạt động offline
- ✅ Hỗ trợ 50+ ngôn ngữ
- ✅ Chính xác cao

**Sau khi cài, khởi động lại ứng dụng.**
        """
    except Exception as e:
        error_msg = str(e)
        yield f"""❌ Lỗi: {error_msg}

💡 **Hướng dẫn khắc phục:**

1. Cài đặt Whisper: `pip install openai-whisper`
2. Cài đặt ffmpeg (nếu chưa có):
   - Windows: `choco install ffmpeg`
   - Hoặc download: https://ffmpeg.org/download.html
3. Đảm bảo file audio hợp lệ (WAV, MP3)
4. Khởi động lại ứng dụng

**Chi tiết lỗi:** {type(e).__name__}
        """


def search_recordings_ui(query):
    """Search recordings by title or notes."""
    if not query or query.strip() == "":
        return "_Nhập từ khóa để tìm kiếm_"
    
    try:
        recordings = audio_manager.get_recordings()
        query_lower = query.lower()
        
        # Filter recordings
        matches = [
            r for r in recordings
            if query_lower in r.get('title', '').lower() or
               query_lower in r.get('notes', '').lower()
        ]
        
        if not matches:
            return f"❌ Không tìm thấy ghi âm nào với từ khóa '{query}'"
        
        # Format results
        output = [f"# 🔍 Tìm thấy {len(matches)} ghi âm\n"]
        
        for i, rec in enumerate(matches, 1):
            output.append(f"## {i}. {rec['title']}")
            output.append(f"**ID:** {rec['id']}")
            output.append(f"**Ngày:** {rec['timestamp'][:19]}")
            output.append(f"**Thời lượng:** {audio_manager.format_duration(rec.get('duration'))}")
            output.append(f"**Trạng thái:** {'✅ Đã xử lý' if rec.get('processed') else '⏳ Chưa xử lý'}")
            if rec.get('notes'):
                output.append(f"**Ghi chú:** {rec['notes']}")
            output.append("\n---\n")
        
        return "\n".join(output)
        
    except Exception as e:
        return f"❌ Lỗi: {str(e)}"


# Semantic Search Functions
def search_meetings_ui(query, meeting_type_filter, language_filter, n_results):
    """Search meetings using semantic search."""
    if not query or query.strip() == "":
        return "⚠️ Vui lòng nhập từ khóa tìm kiếm", ""
    
    if not rag_system:
        return "❌ Hệ thống RAG chưa được khởi tạo", ""

    try:
        # Apply filters
        filter_metadata = {}
        if meeting_type_filter and meeting_type_filter != "Tất cả":
            filter_metadata["meeting_type"] = meeting_type_filter
        if language_filter and language_filter != "Tất cả":
            filter_metadata["language"] = language_filter
        
        results = rag_system.semantic_search(
            query=query,
            k=n_results,
            filter_metadata=filter_metadata if filter_metadata else None
        )
        
        if not results:
            return "❌ Không tìm thấy kết quả phù hợp", ""
        
        # Format results
        markdown_output = []
        markdown_output.append(f"# 🔍 Tìm thấy {len(results)} kết quả\n")
        
        for i, result in enumerate(results, 1):
            metadata = result.get('metadata', {})
            content = result.get('content', '')
            score = result.get('score', 0)
            
            markdown_output.append(f"## {i}. Meeting ID: {metadata.get('meeting_id', 'N/A')}")
            markdown_output.append(f"**Loại:** {metadata.get('meeting_type', 'N/A')} | "
                                  f"**Ngôn ngữ:** {metadata.get('language', 'N/A')} | "
                                  f"**Ngày:** {metadata.get('timestamp', 'N/A')[:10]}")
            
            # Show similarity score if available
            # Note: Score interpretation depends on vector store (distance vs similarity)
            # markdown_output.append(f"**Score:** {score}")
            
            # Show transcript preview
            transcript_preview = content[:300] + "..."
            markdown_output.append(f"\n**Nội dung:**\n> {transcript_preview}\n")
            
            markdown_output.append("---\n")
        
        status = f"✅ Tìm thấy {len(results)} chunks phù hợp với '{query}'"
        
        # Store results for export
        global last_search_results
        last_search_results = results
        
        return status, "\n".join(markdown_output)
        
    except Exception as e:
        return f"❌ Lỗi: {str(e)}", ""


def export_search_results_ui():
    """Export search results to TXT."""
    global last_search_results
    if not last_search_results:
        return None
        
    try:
        filename = f"search_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = Path(filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"SEARCH RESULTS REPORT\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*50 + "\n\n")
            
            for i, result in enumerate(last_search_results, 1):
                metadata = result.get('metadata', {})
                content = result.get('content', '')
                
                f.write(f"RESULT #{i}\n")
                f.write(f"Meeting ID: {metadata.get('meeting_id', 'N/A')}\n")
                f.write(f"Type: {metadata.get('meeting_type', 'N/A')}\n")
                f.write(f"Date: {metadata.get('timestamp', 'N/A')}\n")
                f.write("-" * 20 + "\n")
                f.write(f"{content}\n")
                f.write("="*50 + "\n\n")
                
        return str(filepath)
    except Exception as e:
        return None


def get_vectordb_stats_ui():
    """Get Vector DB statistics."""
    if not rag_system:
        return "❌ Hệ thống RAG chưa được khởi tạo"

    try:
        stats = rag_system.get_stats()
        
        output = []
        output.append(f"# 📊 Thống Kê {stats.get('vector_store', 'Vector DB').upper()}\n")
        output.append(f"**Trạng thái:** {stats.get('status', 'Unknown')}")
        output.append(f"**Tổng số chunks:** {stats.get('total_chunks', 0)}")
        output.append(f"**Embedding Model:** {stats.get('embedding_model', 'N/A')}")
        output.append(f"**LLM Provider:** {stats.get('llm_provider', 'N/A')}")
        
        return "\n".join(output)
        
    except Exception as e:
        return f"❌ Lỗi: {str(e)}"


def get_recordings_list_ui():
    """Get list of recordings."""
    try:
        recordings = audio_manager.get_recordings()
        
        stats = audio_manager.get_statistics()
        stats_text = f"""
📊 **Thống Kê Ghi Âm**
- Tổng số: {stats['total_recordings']}
- Đã xử lý: {stats['processed']}
- Chưa xử lý: {stats['unprocessed']}
- Tổng thời lượng: {stats['total_duration_minutes']:.1f} phút
        """
        
        if not recordings:
            return gr.Dropdown(choices=[], value=None), stats_text
        
        choices = [f"{r['id']} - {r['title']}" for r in recordings]
        
        return gr.Dropdown(choices=choices, value=None), stats_text
    except Exception as e:
        print(f"Error in get_recordings_list_ui: {e}")
        return gr.Dropdown(choices=[], value=None), f"❌ Lỗi: {str(e)}"


def load_recording_info_ui(selected):
    """Load recording information."""
    if not selected:
        return "", "", None
    
    try:
        recording_id = selected.split(" - ")[0]
        recording = audio_manager.get_recording(recording_id)
        
        if not recording:
            return "❌ Không tìm thấy ghi âm", "", None
        
        info = f"""
### 📁 {recording['title']}

**ID:** {recording['id']}  
**Ngày:** {recording['timestamp'][:19]}  
**Thời lượng:** {audio_manager.format_duration(recording.get('duration'))}  
**Trạng thái:** {'✅ Đã xử lý' if recording.get('processed') else '⏳ Chưa xử lý'}  
**Ghi chú:** {recording.get('notes', 'Không có ghi chú')}
        """
        
        return info, recording.get('notes', ''), recording['filepath']
    except Exception as e:
        return f"❌ Lỗi: {str(e)}", "", None


def delete_recording_ui(selected):
    """Delete selected recording."""
    if not selected:
        return "⚠️ Chưa chọn ghi âm", gr.Dropdown(choices=[], value=None), ""
    
    try:
        recording_id = selected.split(" - ")[0]
        success = audio_manager.delete_recording(recording_id)
        
        if success:
            dropdown, stats = get_recordings_list_ui()
            return f"✅ Đã xóa {recording_id}", dropdown, stats
        else:
            return f"❌ Không thể xóa {recording_id}", gr.Dropdown(choices=[], value=None), ""
    except Exception as e:
        return f"❌ Lỗi: {str(e)}", gr.Dropdown(choices=[], value=None), ""


def export_to_docx():
    """Export analysis results to DOCX file."""
    global last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    if not last_summary:
        return None
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading("Meeting Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"File: {current_filename}")
        date_para.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Summary
        doc.add_heading("Summary", 1)
        doc.add_paragraph(last_summary)
        
        # Topics
        if last_topics:
            doc.add_heading("Main Topics", 1)
            for i, topic in enumerate(last_topics, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{topic.get('topic', 'N/A')}").bold = True
                doc.add_paragraph(topic.get('description', ''), style='List Bullet 2')
        
        # Action Items
        if last_actions:
            doc.add_heading("Action Items", 1)
            for i, action in enumerate(last_actions, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{action.get('task', 'N/A')}").bold = True
                doc.add_paragraph(f"Assignee: {action.get('assignee', 'N/A')}", style='List Bullet 2')
                doc.add_paragraph(f"Deadline: {action.get('deadline', 'N/A')}", style='List Bullet 2')
        
        # Decisions
        if last_decisions:
            doc.add_heading("Important Decisions", 1)
            for i, decision in enumerate(last_decisions, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{decision.get('decision', 'N/A')}").bold = True
                doc.add_paragraph(f"Context: {decision.get('context', 'N/A')}", style='List Bullet 2')
        
        filename = f"meeting_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = Path(filename)
        doc.save(str(filepath))
        
        return str(filepath)
    except Exception as e:
        print(f"Error exporting DOCX: {e}")
        return None


# Custom CSS
custom_css = """
.gradio-container {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

.main-header {
    text-align: center;
    padding: 2.5rem;
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    border-radius: 12px;
    margin-bottom: 2rem;
    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
}

.main-header h1 {
    margin: 0;
    font-size: 2.5rem;
    font-weight: 700;
}

.main-header p {
    margin: 0.5rem 0 0 0;
    font-size: 1.1rem;
    opacity: 0.95;
}
"""

# Build UI
with gr.Blocks(css=custom_css, title="Meeting Analyzer Pro", theme=gr.themes.Soft()) as demo:
    
    # Header
    gr.HTML("""
    <div class="main-header">
        <h1>🎯 Meeting Transcript Analyzer Pro</h1>
        <p>Phân tích cuộc họp thông minh với AI - Hỗ trợ đa ngôn ngữ</p>
    </div>
    """)

    
    # Main content in tabs
    with gr.Tabs() as tabs:
        
        # Tab 1: Recording
        with gr.Tab("🎙️ Ghi Âm"):
            with gr.Row():
                with gr.Column(scale=3):
                    with gr.Row():
                        recording_title_input = gr.Textbox(
                            label="Tiêu đề",
                            placeholder=f"Ghi âm {datetime.now().strftime('%d/%m/%Y')}",
                            scale=2
                        )
                        recording_lang_input = gr.Dropdown(
                            label="Ngôn ngữ",
                            choices=[("Tiếng Việt", "vi"), ("English", "en"), ("日本語", "ja"), ("한국어", "ko"), ("中文", "zh")],
                            value="vi",
                            scale=1
                        )
                    
                    audio_recorder_main = gr.Audio(
                        sources=["microphone"],
                        type="filepath",
                        label="🎙️ Ghi âm",
                        waveform_options={"show_recording_waveform": True}
                    )
                    
                    transcript_display = gr.Textbox(
                        label="📝 Transcript",
                        interactive=False,
                        lines=12,
                        placeholder="Nhấn Stop để tự động transcribe bằng OpenAI Whisper..."
                    )
                    
                    with gr.Row():
                        save_recording_btn = gr.Button("💾 Lưu Audio & Transcript", variant="primary", scale=2)
                        clear_recording_btn = gr.Button("🗑️ Hủy", variant="secondary", scale=1)
                    
                    save_status = gr.Textbox(label="Trạng thái", interactive=False, lines=3, show_label=False)
                    recording_id_hidden = gr.Textbox(visible=False)
                
                with gr.Column(scale=1):
                    gr.Markdown("""
                    ### 📝 Cách dùng
                    
                    1. **Chọn ngôn ngữ** ghi âm
                    2. **Click microphone** → Bắt đầu ghi
                    3. **Nhấn Stop** → Tự động:
                       - Lưu file audio
                       - Transcribe bằng Whisper
                    4. **Nhấn Lưu** để lưu vào thư viện
                    5. **Nhấn Hủy** để xóa và ghi lại
                    
                    ---
                    
                    ### ⚙️ Công nghệ
                    
                    - **STT:** OpenAI Whisper
                    - **Định dạng:** WAV
                    - **Hỗ trợ:** 50+ ngôn ngữ
                    
                    ---
                    
                    ### 💡 Lưu ý
                    
                    - Dùng Whisper Local (Offline)
                    - Miễn phí, không cần API key
                    - Cần cài ffmpeg (xem hướng dẫn bên dưới)
                    
                    ---
                    
                    ### 🔧 Cài đặt ffmpeg
                    
                    **Windows:**
                    1. Download: [ffmpeg.org](https://ffmpeg.org/download.html)
                    2. Extract vào `C:\\ffmpeg`
                    3. Thêm `C:\\ffmpeg\\bin` vào PATH
                    4. Khởi động lại terminal
                    
                    **Hoặc dùng Chocolatey:**
                    ```
                    choco install ffmpeg
                    ```
                    """)
        
        # Tab 2: Upload & Analyze
        with gr.Tab("📤 Upload & Phân Tích"):
            with gr.Row():
                with gr.Column(scale=2):
                    gr.Markdown("### 📁 Upload Transcript")
                    file_input = gr.File(
                        label="Chọn file transcript (TXT, DOCX)",
                        file_types=[".txt", ".docx"]
                    )
                    
                    with gr.Row():
                        meeting_type = gr.Dropdown(
                            label="🎯 Loại Cuộc Họp",
                            choices=[
                                ("📋 Meeting - Cuộc họp thông thường", "meeting"),
                                ("🎓 Workshop - Hội thảo/Đào tạo", "workshop"),
                                ("💡 Brainstorming - Động não", "brainstorming")
                            ],
                            value="meeting"
                        )
                        
                        output_lang = gr.Dropdown(
                            label="🌍 Ngôn Ngữ Output",
                            choices=[
                                ("🇻🇳 Tiếng Việt", "vi"),
                                ("🇬🇧 English", "en"),
                                ("🇯🇵 日本語", "ja"),
                                ("🇰🇷 한국어", "ko"),
                                ("🇨🇳 中文", "zh-CN"),
                                ("🇪🇸 Español", "es"),
                                ("🇫🇷 Français", "fr"),
                                ("🇩🇪 Deutsch", "de")
                            ],
                            value="vi"
                        )
                    
                    process_btn = gr.Button(
                        "🚀 Phân Tích Ngay",
                        variant="primary",
                        size="lg"
                    )
                    
                    status_box = gr.Textbox(
                        label="Trạng thái",
                        interactive=False,
                        lines=2
                    )
                
                with gr.Column(scale=1):
                    gr.Markdown("""
                    ### 💡 Hướng Dẫn
                    
                    **Bước 1:** Upload file transcript
                    
                    **Bước 2:** Chọn loại cuộc họp
                    
                    **Bước 3:** Chọn ngôn ngữ output
                    
                    **Bước 4:** Click "Phân Tích Ngay"
                    
                    ---
                    
                    ✨ **Tính năng:**
                    - AI tự động tóm tắt
                    - Trích xuất chủ đề chính
                    - Phát hiện action items
                    - Ghi nhận quyết định quan trọng
                    - Hỗ trợ 8 ngôn ngữ
                    """)
            
            gr.Markdown("---")
            
            # Results section
            gr.Markdown("## 📊 Kết Quả Phân Tích")
            
            with gr.Row():
                with gr.Column():
                    gr.Markdown("### 📝 Tóm Tắt Cuộc Họp")
                    summary_output = gr.Textbox(
                        lines=6,
                        interactive=False,
                        placeholder="Tóm tắt sẽ hiển thị ở đây..."
                    )
            
            with gr.Row():
                with gr.Column():
                    gr.Markdown("### 🎯 Chủ Đề Chính")
                    topics_output = gr.Markdown("_Chưa có dữ liệu_")
                
                with gr.Column():
                    gr.Markdown("### ✅ Action Items")
                    actions_output = gr.Markdown("_Chưa có dữ liệu_")
            
            with gr.Row():
                with gr.Column():
                    gr.Markdown("### 🎯 Quyết Định Quan Trọng")
                    decisions_output = gr.Markdown("_Chưa có dữ liệu_")
        
        # Tab 3: Chat with AI
        with gr.Tab("💬 Chat với AI"):
            gr.Markdown("### 🤖 Hỏi đáp thông minh về cuộc họp")
            
            with gr.Row():
                with gr.Column(scale=3):
                    chatbot_display = gr.Chatbot(
                        height=500,
                        label="Cuộc trò chuyện",
                        type="messages"
                    )
                    
                    with gr.Row():
                        chat_input = gr.Textbox(
                            placeholder="Nhập câu hỏi của bạn...",
                            show_label=False,
                            scale=4
                        )
                        send_btn = gr.Button("📤 Gửi", scale=1, variant="primary")
                    
                    clear_btn = gr.Button("🗑️ Xóa lịch sử", size="sm")
                
                with gr.Column(scale=1):
                    gr.Markdown("### 💡 Câu Hỏi Gợi Ý")
                    
                    q1 = gr.Button("📋 Tóm tắt cuộc họp", size="sm")
                    q2 = gr.Button("👥 Ai tham gia?", size="sm")
                    q3 = gr.Button("✅ Action items là gì?", size="sm")
                    q4 = gr.Button("🎯 Quyết định quan trọng?", size="sm")
                    q5 = gr.Button("📊 Chủ đề chính?", size="sm")
                    
                    gr.Markdown("""
                    ---
                    **💬 Bạn có thể hỏi:**
                    - Tóm tắt cuộc họp
                    - Ai tham gia meeting?
                    - Action items là gì?
                    - Quyết định nào được đưa ra?
                    - Chủ đề chính là gì?
                    """)
        
        # Tab 4: Library
        with gr.Tab("📚 Thư Viện Lưu Trữ"):
            with gr.Tabs():
                # Sub-tab: Analysis History
                with gr.Tab("📊 Lịch Sử Phân Tích"):
                    gr.Markdown("### 📋 Quản lý lịch sử phân tích")
                    
                    with gr.Row():
                        history_dropdown_lib = gr.Dropdown(
                            label="Chọn phân tích",
                            choices=[],
                            interactive=True,
                            scale=3
                        )
                        refresh_history_lib_btn = gr.Button("🔄 Làm mới", scale=1)
                    
                    history_info_lib = gr.Markdown("_Chưa có lịch sử_")
                    
                    with gr.Row():
                        with gr.Column():
                            gr.Markdown("#### 📝 Tóm tắt")
                            summary_lib = gr.Textbox(lines=5, interactive=False)
                        
                        with gr.Column():
                            gr.Markdown("#### 🎯 Chủ đề")
                            topics_lib = gr.Markdown()
                    
                    with gr.Row():
                        load_analysis_btn = gr.Button("📂 Tải vào workspace", variant="primary")
                        delete_analysis_btn = gr.Button("🗑️ Xóa", variant="stop")
                
                # Sub-tab: Recording History
                with gr.Tab("🎙️ Lịch Sử Ghi Âm"):
                    gr.Markdown("### 🎤 Quản lý ghi âm")
                    
                    with gr.Row():
                        recordings_dropdown = gr.Dropdown(
                            label="Chọn ghi âm",
                            choices=[],
                            interactive=True,
                            scale=3
                        )
                        refresh_recordings_btn = gr.Button("🔄 Làm mới", scale=1)
                    
                    recordings_stats = gr.Markdown("📊 Chưa có ghi âm nào")
                    
                    with gr.Row():
                        with gr.Column(scale=2):
                            recording_info_display = gr.Markdown("_Chọn ghi âm để xem chi tiết_")
                            recording_notes_display = gr.Textbox(label="Ghi chú", interactive=False, lines=3)
                        
                        with gr.Column(scale=1):
                            audio_player = gr.Audio(
                                label="Phát lại",
                                interactive=False
                            )
                            
                            with gr.Row():
                                process_recording_btn = gr.Button("🔄 Xử lý", variant="primary")
                                delete_recording_btn = gr.Button("🗑️ Xóa", variant="stop")
                            
                            delete_status = gr.Textbox(label="Trạng thái", interactive=False)
        
        # Tab 5: Semantic Search
        with gr.Tab("🔍 Tìm Kiếm Thông Minh"):
            with gr.Tabs():
                # Sub-tab: Search Analysis
                with gr.Tab("📊 Tìm Phân Tích"):
                    gr.Markdown("### 🔍 Tìm kiếm lịch sử phân tích")
                    
                    # Statistics
                    with gr.Accordion("📊 Thống Kê Database", open=False):
                        stats_display = gr.Markdown()
                        refresh_stats_btn = gr.Button("🔄 Làm mới thống kê")
                    
                    # Search interface
                    with gr.Row():
                        with gr.Column(scale=3):
                            search_query = gr.Textbox(
                                label="Tìm kiếm",
                                placeholder="VD: React Hooks training, budget planning, team meeting...",
                                lines=2
                            )
                        with gr.Column(scale=1):
                            meeting_type_filter = gr.Dropdown(
                                label="Loại Meeting",
                                choices=["Tất cả", "meeting", "workshop", "brainstorming"],
                                value="Tất cả"
                            )
                            language_filter = gr.Dropdown(
                                label="Ngôn ngữ",
                                choices=["Tất cả", "en", "vi", "ja", "ko", "zh-CN"],
                                value="Tất cả"
                            )
                            n_results = gr.Slider(
                                label="Số kết quả",
                                minimum=1,
                                maximum=10,
                                value=5,
                                step=1
                            )
                    
                    search_btn = gr.Button("🔍 Tìm kiếm", variant="primary", size="lg")
                    search_status = gr.Textbox(label="Trạng thái", interactive=False)
                    search_results = gr.Markdown()
                    
                    with gr.Row():
                        export_search_btn = gr.Button("💾 Xuất Kết Quả (TXT)", size="sm")
                        search_export_file = gr.File(label="File kết quả")
                    
                    # Examples
                    gr.Examples(
                        examples=[
                            ["React Hooks training", "workshop", "en", 3],
                            ["budget planning meeting", "meeting", "Tất cả", 5],
                            ["brainstorming new features", "brainstorming", "Tất cả", 3],
                        ],
                        inputs=[search_query, meeting_type_filter, language_filter, n_results]
                    )
                
                # Sub-tab: Search Recordings
                with gr.Tab("🎙️ Tìm Ghi Âm"):
                    gr.Markdown("### 🔍 Tìm kiếm ghi âm cuộc gọi")
                    
                    with gr.Row():
                        search_recording_query = gr.Textbox(
                            label="Tìm kiếm theo tiêu đề hoặc ghi chú",
                            placeholder="VD: Team meeting, Sprint planning...",
                            scale=3
                        )
                        search_recording_btn = gr.Button("🔍 Tìm", variant="primary", scale=1)
                    
                    search_recording_results = gr.Markdown("_Nhập từ khóa để tìm kiếm_")
                    
                    gr.Markdown("""
                    ---
                    **💡 Mẹo tìm kiếm:**
                    - Tìm theo tiêu đề: "Team meeting"
                    - Tìm theo ngày: "25/11/2024"
                    - Tìm theo ghi chú: "Sprint planning"
                    """)
        
        # Tab 6: Export Results
        with gr.Tab("📄 Xuất Kết Quả"):
            gr.Markdown("### 📥 Tải xuống kết quả phân tích")
            
            with gr.Row():
                with gr.Column():
                    gr.Markdown("""
                    #### 📄 Xuất file TXT
                    - Format đơn giản, dễ đọc
                    - Phù hợp để chia sẻ qua email
                    - Mở được trên mọi thiết bị
                    """)
                    export_txt_btn = gr.Button("📄 Xuất TXT", size="lg", variant="primary")
                
                with gr.Column():
                    gr.Markdown("""
                    #### 📝 Xuất file DOCX
                    - Format chuyên nghiệp
                    - Có thể chỉnh sửa trong Word
                    - Phù hợp cho báo cáo
                    """)
                    export_docx_btn = gr.Button("📝 Xuất DOCX", size="lg", variant="primary")
            
            export_file = gr.File(label="File đã xuất")
            
            gr.Markdown("""
            ---
            **💡 Lưu ý:**
            - Cần xử lý transcript trước khi xuất
            - File sẽ được lưu với timestamp
            - Hỗ trợ đa ngôn ngữ
            """)
    
    # Footer
    gr.Markdown("""
    ---
    <div style="text-align: center; color: #666; padding: 1rem;">
        <p>🚀 Powered by Gemini AI | �️ Audio aRecording | 🔍 Semantic Search | 🌍 Multi-language</p>
        <p style="font-size: 0.9em;">Version 3.1 - Sprint 3 Complete Edition</p>
    </div>
    """)
    
    # Event handlers
    
    # Recording Tab handlers
    
    # Auto-transcribe when recording stops (audio changes)
    audio_recorder_main.stop_recording(
        fn=transcribe_audio_whisper,
        inputs=[audio_recorder_main, recording_lang_input],
        outputs=[transcript_display]
    )
    
    save_recording_btn.click(
        fn=save_recording_and_transcribe,
        inputs=[audio_recorder_main, recording_title_input, recording_lang_input, transcript_display],
        outputs=[save_status, recording_id_hidden]
    )
    
    clear_recording_btn.click(
        fn=lambda: (None, "", "", ""),
        outputs=[audio_recorder_main, transcript_display, save_status, recording_id_hidden]
    )
    
    # Library - Analysis History handlers
    refresh_history_lib_btn.click(
        fn=refresh_history,
        outputs=[history_dropdown_lib, history_info_lib]
    )
    
    load_analysis_btn.click(
        fn=load_history,
        inputs=[history_dropdown_lib],
        outputs=[status_box, summary_output, topics_output, actions_output, decisions_output]
    )
    
    # Auto-refresh on page load
    demo.load(
        fn=get_vectordb_stats_ui,
        outputs=[stats_display]
    ).then(
        fn=get_recordings_list_ui,
        outputs=[recordings_dropdown, recordings_stats]
    ).then(
        fn=refresh_history,
        outputs=[history_dropdown_lib, history_info_lib]
    )
    
    # Process handler
    process_btn.click(
        fn=process_file,
        inputs=[file_input, meeting_type, output_lang],
        outputs=[status_box, summary_output, topics_output, actions_output, decisions_output]
    ).then(
        fn=refresh_history,  # Refresh history after processing
        outputs=[history_dropdown_lib, history_info_lib]
    )
    
    # Export handlers
    export_txt_btn.click(
        fn=export_to_txt,
        outputs=[export_file]
    )
    
    export_docx_btn.click(
        fn=export_to_docx,
        outputs=[export_file]
    )
    
    # Recording Library handlers
    refresh_recordings_btn.click(
        fn=get_recordings_list_ui,
        outputs=[recordings_dropdown, recordings_stats]
    )
    
    recordings_dropdown.change(
        fn=load_recording_info_ui,
        inputs=[recordings_dropdown],
        outputs=[recording_info_display, recording_notes_display, audio_player]
    )
    
    delete_recording_btn.click(
        fn=delete_recording_ui,
        inputs=[recordings_dropdown],
        outputs=[delete_status, recordings_dropdown, recordings_stats]
    )
    
    # Semantic Search handlers
    refresh_stats_btn.click(
        fn=get_vectordb_stats_ui,
        outputs=[stats_display]
    )
    
    search_btn.click(
        fn=search_meetings_ui,
        inputs=[search_query, meeting_type_filter, language_filter, n_results],
        outputs=[search_status, search_results]
    )
    
    search_recording_btn.click(
        fn=search_recordings_ui,
        inputs=[search_recording_query],
        outputs=[search_recording_results]
    )
    
    export_search_btn.click(
        fn=export_search_results_ui,
        outputs=[search_export_file]
    )
    
    # Chat handlers
    send_btn.click(
        fn=chat_with_ai,
        inputs=[chat_input, chatbot_display],
        outputs=[chatbot_display]
    ).then(
        fn=lambda: "",
        outputs=[chat_input]
    )
    
    chat_input.submit(
        fn=chat_with_ai,
        inputs=[chat_input, chatbot_display],
        outputs=[chatbot_display]
    ).then(
        fn=lambda: "",
        outputs=[chat_input]
    )
    
    clear_btn.click(
        fn=lambda: [],
        outputs=[chatbot_display]
    )
    
    # Quick questions
    def send_q(q, h):
        if h is None:
            h = []
        return chat_with_ai(q, h)
    
    q1.click(fn=lambda h: send_q("Tóm tắt cuộc họp này", h), inputs=[chatbot_display], outputs=[chatbot_display])
    q2.click(fn=lambda h: send_q("Ai tham gia cuộc họp?", h), inputs=[chatbot_display], outputs=[chatbot_display])
    q3.click(fn=lambda h: send_q("Action items là gì?", h), inputs=[chatbot_display], outputs=[chatbot_display])
    q4.click(fn=lambda h: send_q("Những quyết định quan trọng nào được đưa ra?", h), inputs=[chatbot_display], outputs=[chatbot_display])
    q5.click(fn=lambda h: send_q("Chủ đề chính của cuộc họp là gì?", h), inputs=[chatbot_display], outputs=[chatbot_display])


if __name__ == "__main__":
    demo.launch(
        server_name="localhost",
        server_port=7777,
        share=False
    )


def process_upload(file_type, audio_file, text_file, transcribe_lang, enable_diarization, meeting_type, output_lang):
    """Process uploaded file - audio or text with optional speaker diarization.
    
    Args:
        file_type: "audio" or "text"
        audio_file: Audio file path (if file_type == "audio")
        text_file: Text file path (if file_type == "text")
        transcribe_lang: Language for audio transcription
        enable_diarization: Whether to enable speaker diarization
        meeting_type: Type of meeting
        output_lang: Output language for analysis
        
    Returns:
        Tuple of (status, transcript, summary, topics, actions, decisions, participants)
    """
    global chatbot, transcript_text, last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    logger.info(f"Processing upload: type={file_type}, meeting_type={meeting_type}, diarization={enable_diarization}")
    
    try:
        # Step 1: Get transcript
        transcript = ""
        status_msg = ""
        
        if file_type == "audio":
            if audio_file is None:
                return "❌ Vui lòng upload file audio!", "", "", "", "", "", ""
            
            # Validate audio file
            is_valid, error_msg = validate_file(
                audio_file.name if hasattr(audio_file, 'name') else audio_file,
                max_size_mb=200,
                allowed_extensions=['.wav', '.mp3', '.m4a', '.webm', '.ogg']
            )
            if not is_valid:
                return f"❌ {error_msg}", "", "", "", "", "", ""
            
            status_msg = "🔄 Đang transcribe audio..."
            
            # Transcribe with or without speaker diarization
            if enable_diarization:
                # Use speaker diarization
                from src.audio.speaker_diarization import transcribe_with_speakers
                
                status_msg = "🔄 Đang transcribe với speaker diarization..."
                
                # Get transcript with speakers
                for message in transcribe_with_speakers(audio_file, transcribe_lang):
                    if message.startswith("📝 **Complete Transcript"):
                        # Extract transcript from final message
                        lines = message.split('\n')
                        transcript_lines = []
                        in_transcript = False
                        for line in lines:
                            if line.startswith('[') and ']:' in line:
                                in_transcript = True
                                transcript_lines.append(line)
                            elif in_transcript and line.strip() and not line.startswith('---'):
                                transcript_lines.append(line)
                        transcript = '\n'.join(transcript_lines)
                        break
                
                if not transcript:
                    return "❌ Không thể transcribe audio với speaker diarization", "", "", "", "", "", ""
                
            else:
                # Regular transcription without diarization
                from src.audio.huggingface_stt import transcribe_audio_huggingface
                
                for message in transcribe_audio_huggingface(audio_file, transcribe_lang, realtime=False):
                    if message.startswith("📝 **Transcript"):
                        # Extract transcript
                        lines = message.split('\n')
                        transcript_lines = []
                        for i, line in enumerate(lines):
                            if i > 0 and line.strip() and not line.startswith('✅') and not line.startswith('💡'):
                                transcript_lines.append(line)
                        transcript = '\n'.join(transcript_lines).strip()
                        break
                
                if not transcript:
                    return "❌ Không thể transcribe audio", "", "", "", "", "", ""
            
            current_filename = Path(audio_file.name if hasattr(audio_file, 'name') else audio_file).name
            
        else:  # text file
            if text_file is None:
                return "❌ Vui lòng upload file text!", "", "", "", "", "", ""
            
            # Validate text file
            is_valid, error_msg = validate_file(
                text_file.name if hasattr(text_file, 'name') else text_file,
                max_size_mb=50,
                allowed_extensions=['.txt', '.docx']
            )
            if not is_valid:
                return f"❌ {error_msg}", "", "", "", "", "", ""
            
            # Load text file
            loader = TranscriptLoader()
            transcript = loader.load_file(text_file.name if hasattr(text_file, 'name') else text_file)
            current_filename = Path(text_file.name if hasattr(text_file, 'name') else text_file).name
        
        # Sanitize and preprocess transcript
        transcript = sanitize_input(transcript, max_length=50000)
        preprocessor = TranscriptPreprocessor()
        transcript = preprocessor.clean_text(transcript)
        transcript = preprocessor.truncate_text(transcript, max_length=15000)
        transcript_text = transcript
        
        logger.info(f"Transcript ready: {len(transcript)} chars")
        
        # Step 2: Analyze with LLM
        status_msg = "🔄 Đang phân tích với AI..."
        
        # Import meeting type functions
        from src.rag.meeting_types import WorkshopFunctions, BrainstormingFunctions
        
        # Initialize LLM
        llm_manager = LLMManager(
            provider="gemini",
            model_name="gemini-2.5-flash",
            temperature=Settings.TEMPERATURE,
            max_tokens=Settings.MAX_TOKENS,
            api_key=Settings.GEMINI_API_KEY,
        )
        
        chatbot = Chatbot(
            llm_manager=llm_manager,
            transcript=transcript,
            language=output_lang,
            meeting_type=meeting_type
        )
        
        # Generate analysis
        summary = chatbot.generate_summary()
        topics = chatbot.extract_topics()
        action_items = chatbot.extract_action_items_initially()
        decisions = chatbot.extract_decisions()
        
        # Extract participants (if speaker diarization was used)
        participants_text = "_Không có thông tin_"
        if enable_diarization and file_type == "audio":
            # Extract unique speakers from transcript
            import re
            speakers = set()
            for line in transcript.split('\n'):
                match = re.match(r'\[\d+:\d+\] \*\*(Guest-\d+)\*\*:', line)
                if match:
                    speakers.add(match.group(1))
            
            if speakers:
                participants_text = f"**Số người tham gia:** {len(speakers)}\n\n"
                participants_text += "**Danh sách:**\n"
                for speaker in sorted(speakers):
                    participants_text += f"- {speaker}\n"
        
        # Extract specialized data
        specialized_data = {}
        if meeting_type == "workshop":
            specialized_data['key_learnings'] = WorkshopFunctions.extract_key_learnings(transcript)
            specialized_data['exercises'] = WorkshopFunctions.extract_exercises(transcript)
            specialized_data['qa_pairs'] = WorkshopFunctions.extract_qa_pairs(transcript)
        elif meeting_type == "brainstorming":
            ideas_result = BrainstormingFunctions.extract_ideas(transcript)
            specialized_data['ideas'] = ideas_result
            specialized_data['categorized_ideas'] = BrainstormingFunctions.categorize_ideas(ideas_result)
            specialized_data['concerns'] = BrainstormingFunctions.extract_concerns(transcript)
        
        # Save results
        last_summary = summary
        last_topics = topics
        last_actions = action_items
        last_decisions = decisions
        current_language = output_lang
        
        # Save to history
        try:
            history_manager.save_analysis(
                filename=current_filename,
                summary=summary,
                topics=topics,
                action_items=action_items,
                decisions=decisions,
                metadata={
                    "language": output_lang,
                    "meeting_type": meeting_type,
                    "file_type": file_type,
                    "speaker_diarization": enable_diarization,
                    "specialized_data": specialized_data
                }
            )
        except Exception as e:
            logger.warning(f"Failed to save history: {e}")
        
        # Add to RAG system
        if rag_system:
            try:
                rag_metadata = {
                    "meeting_type": meeting_type,
                    "language": output_lang,
                    "filename": current_filename,
                    "timestamp": datetime.now().isoformat()
                }
                # Use filename as meeting_id for simplicity, or generate a UUID
                meeting_id = f"{Path(current_filename).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                rag_system.add_meeting(meeting_id, transcript, rag_metadata)
                logger.info(f"Added meeting {meeting_id} to RAG system")
            except Exception as e:
                logger.error(f"Failed to add to RAG: {e}")
        
        # Format outputs
        topics_text = format_topics_by_type(topics, meeting_type, specialized_data, output_lang)
        actions_text = format_actions(action_items, output_lang)
        decisions_text = format_decisions_by_type(decisions, meeting_type, specialized_data, output_lang)
        
        success_msg = f"✅ Hoàn thành: {current_filename} | Loại: {meeting_type}"
        if enable_diarization:
            success_msg += " | Speaker Diarization: ON"
        
        return (
            success_msg,
            transcript,
            summary,
            topics_text,
            actions_text,
            decisions_text,
            participants_text
        )
        
    except Exception as e:
        logger.error(f"Error processing upload: {str(e)}", exc_info=True)
        user_message = get_user_friendly_message(e, output_lang)
        return f"❌ Lỗi: {user_message}", "", "", "", "", "", ""
