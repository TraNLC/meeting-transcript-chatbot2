"""Gradio UI for Meeting Transcript Chatbot."""

import gradio as gr
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add src to path
sys.path.append(str(Path(__file__).parent.parent.parent))

from src.config import Settings
from src.data import TranscriptLoader, TranscriptPreprocessor
from src.data.history_manager import HistoryManager
from src.llm import LLMManager
from src.rag import Chatbot


# Global variables
chatbot = None
transcript_text = ""
last_summary = ""
last_topics = []
last_actions = []
last_decisions = []
current_language = "vi"
current_filename = ""
history_manager = HistoryManager()


def process_file(file):
    """Process uploaded transcript file."""
    global chatbot, transcript_text, last_summary, last_topics, last_actions, last_decisions, current_language, current_filename
    
    # Fixed configuration
    provider = "gemini"
    model = "gemini-2.5-flash"
    language = "vi"
    
    upload_msg = {
        "vi": "❌ Vui lòng upload file!",
        "en": "❌ Please upload a file!",
        "ja": "❌ ファイルをアップロードしてください！",
        "ko": "❌ 파일을 업로드해주세요!",
        "zh": "❌ 请上传文件！"
    }
    
    if file is None:
        return upload_msg.get(language, upload_msg["vi"]), "", "", "", ""
    
    # Save current language
    current_language = language
    
    try:
        # Load transcript
        loader = TranscriptLoader()
        transcript = loader.load_file(file.name)
        
        # Clean and truncate
        preprocessor = TranscriptPreprocessor()
        transcript = preprocessor.clean_text(transcript)
        transcript = preprocessor.truncate_text(transcript, max_length=15000)
        transcript_text = transcript
        
        # Initialize LLM and chatbot with Gemini
        llm_manager = LLMManager(
            provider=provider,
            model_name=model,
            temperature=Settings.TEMPERATURE,
            max_tokens=Settings.MAX_TOKENS,
            api_key=Settings.GEMINI_API_KEY,
        )
        
        chatbot = Chatbot(llm_manager=llm_manager, transcript=transcript, language=language)
        
        # Generate summary
        summary = chatbot.generate_summary()
        
        # Extract information
        topics = chatbot.extract_topics()
        # Use only once, then cache for later use
        action_items = chatbot.extract_action_items()
        # After this, always use last_actions for action items
        decisions = chatbot.extract_decisions()
        
        # Save results globally for export
        last_summary = summary
        last_topics = topics
        last_actions = action_items
        last_decisions = decisions
        current_filename = file.name
        
        # Save to history
        try:
            history_id = history_manager.save_analysis(
                filename=file.name,
                summary=summary,
                topics=topics,
                action_items=action_items,
                decisions=decisions,
                metadata={"language": language, "provider": provider, "model": model}
            )
            print(f"✅ Saved to history: {history_id}")
        except Exception as e:
            print(f"⚠️ Failed to save history: {e}")
        
        # Format outputs
        topics_text = format_topics(topics, language)
        actions_text = format_actions(action_items, language)
        decisions_text = format_decisions(decisions, language)
        
        success_msgs = {
            "vi": "✅ Transcript đã được xử lý thành công!",
            "en": "✅ Transcript processed successfully!",
            "ja": "✅ トランスクリプトが正常に処理されました！",
            "ko": "✅ 트랜스크립트가 성공적으로 처리되었습니다!",
            "zh": "✅ 记录处理成功！"
        }
        
        return (
            success_msgs.get(language, success_msgs["vi"]),
            summary,
            topics_text,
            actions_text,
            decisions_text
        )
        
    except Exception as e:
        error_prefixes = {
            "vi": "❌ Lỗi:",
            "en": "❌ Error:",
            "ja": "❌ エラー:",
            "ko": "❌ 오류:",
            "zh": "❌ 错误:"
        }
        return f"{error_prefixes.get(language, error_prefixes['vi'])} {str(e)}", "", "", "", ""


def format_topics(topics, language="vi"):
    """Format topics for display."""
    no_topics_msg = {
        "vi": "Không tìm thấy chủ đề rõ ràng",
        "en": "No clear topics found",
        "ja": "明確なトピックが見つかりません",
        "ko": "명확한 주제를 찾을 수 없습니다",
        "zh": "未找到明确的主题"
    }
    no_desc_msg = {
        "vi": "Không có mô tả",
        "en": "No description",
        "ja": "説明なし",
        "ko": "설명 없음",
        "zh": "无描述"
    }
    
    if not topics:
        return no_topics_msg.get(language, no_topics_msg["vi"])
    
    result = []
    for i, topic in enumerate(topics, 1):
        result.append(f"**{i}. {topic.get('topic', 'N/A')}**")
        result.append(f"   {topic.get('description', no_desc_msg.get(language, no_desc_msg['vi']))}")
        result.append("")
    
    return "\n".join(result)


def format_actions(actions, language="vi"):
    """Format action items for display."""
    no_actions_msg = {
        "vi": "Không tìm thấy action items",
        "en": "No action items found",
        "ja": "アクションアイテムが見つかりません",
        "ko": "액션 아이템을 찾을 수 없습니다",
        "zh": "未找到行动项"
    }
    assignee_labels = {
        "vi": "Người phụ trách",
        "en": "Assignee",
        "ja": "担当者",
        "ko": "담당자",
        "zh": "负责人"
    }
    not_assigned_msg = {
        "vi": "Chưa phân công",
        "en": "Not assigned",
        "ja": "未割り当て",
        "ko": "미할당",
        "zh": "未分配"
    }
    not_specified_msg = {
        "vi": "Chưa xác định",
        "en": "Not specified",
        "ja": "未指定",
        "ko": "미지정",
        "zh": "未指定"
    }
    
    if not actions:
        return no_actions_msg.get(language, no_actions_msg["vi"])
    
    result = []
    for i, item in enumerate(actions, 1):
        result.append(f"**{i}. {item.get('task', 'N/A')}**")
        result.append(f"   👤 {assignee_labels.get(language, assignee_labels['vi'])}: {item.get('assignee', not_assigned_msg.get(language, not_assigned_msg['vi']))}")
        result.append(f"   📅 Deadline: {item.get('deadline', not_specified_msg.get(language, not_specified_msg['vi']))}")
        result.append("")
    
    return "\n".join(result)


def format_decisions(decisions, language="vi"):
    """Format decisions for display."""
    no_decisions_msg = {
        "vi": "Không tìm thấy quyết định",
        "en": "No decisions found",
        "ja": "決定が見つかりません",
        "ko": "결정을 찾을 수 없습니다",
        "zh": "未找到决定"
    }
    context_labels = {
        "vi": "Bối cảnh",
        "en": "Context",
        "ja": "文脈",
        "ko": "맥락",
        "zh": "背景"
    }
    
    if not decisions:
        return no_decisions_msg.get(language, no_decisions_msg["vi"])
    
    result = []
    for i, decision in enumerate(decisions, 1):
        result.append(f"**{i}. {decision.get('decision', 'N/A')}**")
        result.append(f"   📝 {context_labels.get(language, context_labels['vi'])}: {decision.get('context', 'N/A')}")
        result.append("")
    
    return "\n".join(result)


def export_to_txt():
    """Export analysis results to TXT file."""
    global last_summary, last_topics, last_actions, last_decisions, current_language
    
    if not last_summary:
        return None
    
    # Get labels based on language
    labels = {
        "vi": {
            "title": "BÁO CÁO PHÂN TÍCH CUỘC HỌP",
            "summary": "TÓM TẮT CUỘC HỌP",
            "topics": "CHỦ ĐỀ CHÍNH",
            "actions": "ACTION ITEMS",
            "decisions": "QUYẾT ĐỊNH QUAN TRỌNG",
            "task": "Nhiệm vụ",
            "assignee": "Người phụ trách",
            "deadline": "Deadline",
            "decision": "Quyết định",
            "context": "Bối cảnh"
        },
        "en": {
            "title": "MEETING ANALYSIS REPORT",
            "summary": "MEETING SUMMARY",
            "topics": "MAIN TOPICS",
            "actions": "ACTION ITEMS",
            "decisions": "IMPORTANT DECISIONS",
            "task": "Task",
            "assignee": "Assignee",
            "deadline": "Deadline",
            "decision": "Decision",
            "context": "Context"
        },
        "ja": {
            "title": "会議分析レポート",
            "summary": "会議要約",
            "topics": "主要トピック",
            "actions": "アクションアイテム",
            "decisions": "重要な決定",
            "task": "タスク",
            "assignee": "担当者",
            "deadline": "期限",
            "decision": "決定",
            "context": "文脈"
        },
        "ko": {
            "title": "회의 분석 보고서",
            "summary": "회의 요약",
            "topics": "주요 주제",
            "actions": "액션 아이템",
            "decisions": "중요한 결정",
            "task": "작업",
            "assignee": "담당자",
            "deadline": "마감일",
            "decision": "결정",
            "context": "맥락"
        },
        "zh": {
            "title": "会议分析报告",
            "summary": "会议摘要",
            "topics": "主要主题",
            "actions": "行动项",
            "decisions": "重要决定",
            "task": "任务",
            "assignee": "负责人",
            "deadline": "截止日期",
            "decision": "决定",
            "context": "背景"
        }
    }
    
    lang = labels.get(current_language, labels["vi"])
    
    # Create content
    content = []
    content.append("=" * 80)
    content.append(lang["title"].center(80))
    content.append("=" * 80)
    content.append(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    content.append("=" * 80)
    
    # Summary
    content.append(f"\n{lang['summary']}")
    content.append("-" * 80)
    content.append(last_summary)
    
    # Topics
    if last_topics:
        content.append(f"\n\n{lang['topics']}")
        content.append("-" * 80)
        for i, topic in enumerate(last_topics, 1):
            content.append(f"\n{i}. {topic.get('topic', 'N/A')}")
            content.append(f"   {topic.get('description', '')}")
    
    # Action Items
    if last_actions:
        content.append(f"\n\n{lang['actions']}")
        content.append("-" * 80)
        for i, action in enumerate(last_actions, 1):
            content.append(f"\n{i}. {lang['task']}: {action.get('task', 'N/A')}")
            content.append(f"   {lang['assignee']}: {action.get('assignee', 'N/A')}")
            content.append(f"   {lang['deadline']}: {action.get('deadline', 'N/A')}")
    
    # Decisions
    if last_decisions:
        content.append(f"\n\n{lang['decisions']}")
        content.append("-" * 80)
        for i, decision in enumerate(last_decisions, 1):
            content.append(f"\n{i}. {lang['decision']}: {decision.get('decision', 'N/A')}")
            content.append(f"   {lang['context']}: {decision.get('context', 'N/A')}")
    
    content.append("\n\n" + "=" * 80)
    
    # Save to file
    filename = f"meeting_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    filepath = Path(filename)
    filepath.write_text("\n".join(content), encoding='utf-8')
    
    return str(filepath)


def chat_with_ai(message, history):
    """Chat with AI using function calling.
    
    Args:
        message: User message
        history: Chat history
        
    Returns:
        Updated history
    """
    global chatbot, transcript_text, current_language
    
    if not transcript_text:
        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": "⚠️ Vui lòng upload và xử lý transcript trước khi chat!"})
        return history
    
    try:
        from src.llm.prompts import FunctionCallingSchemas
        # Use last_actions for action items
        
        # Detect if user wants specific function
        message_lower = message.lower()
        
        # Simple rule-based function calling (Sprint 2 demo)
        # In production, use LLM to decide which function to call
        
        response = ""
        function_called = None
        
        if any(keyword in message_lower for keyword in ["task", "action", "nhiệm vụ", "việc"]):
            # Use cached action items (last_actions)
            import re
            global last_actions
            name_match = re.search(r'\b([A-Z][a-z]+)\b', message)
            assignee = name_match.group(1) if name_match else None
            if assignee:
                items = [item for item in last_actions if item.get("assignee", "").lower() == assignee.lower()]
            else:
                items = last_actions
            if items:
                response = f"🔍 Tìm thấy {len(items)} action items:\n\n"
                for i, item in enumerate(items, 1):
                    response += f"{i}. **{item['task']}**\n"
                    response += f"   👤 {item['assignee']} | 📅 {item['deadline']}\n\n"
            else:
                response = "Không tìm thấy action items phù hợp."
        
        elif any(keyword in message_lower for keyword in ["tìm", "search", "keyword"]):
            # Search transcript
            import re
            # Extract keyword from quotes or after "tìm"
            keyword_match = re.search(r'["\']([^"\']+)["\']|tìm\s+(\w+)|search\s+(\w+)', message_lower)
            keyword = None
            if keyword_match:
                keyword = keyword_match.group(1) or keyword_match.group(2) or keyword_match.group(3)
            
            if keyword:
                result = executor.execute("search_transcript", {"keyword": keyword, "context_lines": 2})
                function_called = "search_transcript"
                
                import json
                data = json.loads(result)
                total = data.get("total_matches", 0)
                results = data.get("results", [])
                
                if total > 0:
                    response = f"🔍 Tìm thấy '{keyword}' {total} lần:\n\n"
                    for i, match in enumerate(results[:3], 1):  # Show first 3
                        response += f"**{i}. Dòng {match['line_number']}:**\n"
                        response += f"```\n{match['context']}\n```\n\n"
                    
                    if total > 3:
                        response += f"_... và {total - 3} kết quả khác_"
                else:
                    response = f"Không tìm thấy '{keyword}' trong transcript."
            else:
                response = "Vui lòng cung cấp từ khóa cần tìm. Ví dụ: 'Tìm \"budget\"'"
        
        elif any(keyword in message_lower for keyword in ["người", "participant", "tham gia", "ai"]):
            # Get participants
            result = executor.execute("get_meeting_participants", {})
            function_called = "get_meeting_participants"
            
            import json
            data = json.loads(result)
            participants = data.get("participants", [])
            
            if participants:
                response = f"👥 Có {len(participants)} người tham gia:\n\n"
                for i, p in enumerate(participants, 1):
                    response += f"{i}. **{p['name']}** - {p['role']}\n"
            else:
                response = "Không xác định được người tham gia."
        
        elif any(keyword in message_lower for keyword in ["quyết định", "decision", "kết luận"]):
            # Extract decisions
            result = executor.execute("extract_decisions", {})
            function_called = "extract_decisions"
            
            import json
            data = json.loads(result)
            decisions = data.get("decisions", [])
            
            if decisions:
                response = f"📋 Tìm thấy {len(decisions)} quyết định:\n\n"
                for i, d in enumerate(decisions, 1):
                    response += f"{i}. **{d['decision']}**\n"
                    response += f"   📝 {d['context']}\n\n"
            else:
                response = "Không tìm thấy quyết định rõ ràng."
        
        else:
            # General Q&A using chatbot
            result = chatbot.ask_question(message)
            response = result.get("answer", "Xin lỗi, tôi không hiểu câu hỏi.")
            function_called = "general_qa"
        
        # Add function call info
        if function_called and function_called != "general_qa":
            response = f"🔧 _Function called: `{function_called}`_\n\n{response}"
        
        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": response})
        
    except Exception as e:
        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": f"❌ Lỗi: {str(e)}"})
    return history


def clear_chat():
    """Clear chat history."""
    return []


def refresh_history():
    """Refresh history dropdown."""
    history_list = history_manager.list_history(limit=20)
    
    if not history_list:
        return gr.Dropdown(choices=[], value=None), "_Chưa có lịch sử_"
    
    choices = [
        (f"{item['timestamp'][:10]} - {item['original_file']}", item['id'])
        for item in history_list
    ]
    
    info = f"📊 Tìm thấy {len(history_list)} phân tích đã lưu"
    
    return gr.Dropdown(choices=choices, value=None), info


def load_history(history_id):
    """Load analysis from history."""
    global last_summary, last_topics, last_actions, last_decisions, current_language, transcript_text
    
    if not history_id:
        return "⚠️ Vui lòng chọn phân tích từ danh sách", "", "", "", ""
    
    data = history_manager.load_analysis(history_id)
    
    if not data:
        return "❌ Không tìm thấy phân tích", "", "", "", ""
    
    # Load data
    last_summary = data.get("summary", "")
    last_topics = data.get("topics", [])
    last_actions = data.get("action_items", [])
    last_decisions = data.get("decisions", [])
    current_language = data.get("metadata", {}).get("language", "vi")
    
    # Format outputs
    topics_text = format_topics(last_topics, current_language)
    actions_text = format_actions(last_actions, current_language)
    decisions_text = format_decisions(last_decisions, current_language)
    
    status = f"✅ Đã tải phân tích: {data.get('original_file')} ({data.get('timestamp')[:10]})"
    
    return status, last_summary, topics_text, actions_text, decisions_text


def export_to_docx():
    """Export analysis results to DOCX file."""
    global last_summary, last_topics, last_actions, last_decisions, current_language
    
    if not last_summary:
        return None
    
    # Get labels based on language
    labels = {
        "vi": {
            "title": "BÁO CÁO PHÂN TÍCH CUỘC HỌP",
            "summary": "TÓM TẮT CUỘC HỌP",
            "topics": "CHỦ ĐỀ CHÍNH",
            "actions": "ACTION ITEMS",
            "decisions": "QUYẾT ĐỊNH QUAN TRỌNG",
            "task": "Nhiệm vụ",
            "assignee": "Người phụ trách",
            "deadline": "Deadline",
            "decision": "Quyết định",
            "context": "Bối cảnh"
        },
        "en": {
            "title": "MEETING ANALYSIS REPORT",
            "summary": "MEETING SUMMARY",
            "topics": "MAIN TOPICS",
            "actions": "ACTION ITEMS",
            "decisions": "IMPORTANT DECISIONS",
            "task": "Task",
            "assignee": "Assignee",
            "deadline": "Deadline",
            "decision": "Decision",
            "context": "Context"
        },
        "ja": {
            "title": "会議分析レポート",
            "summary": "会議要約",
            "topics": "主要トピック",
            "actions": "アクションアイテム",
            "decisions": "重要な決定",
            "task": "タスク",
            "assignee": "担当者",
            "deadline": "期限",
            "decision": "決定",
            "context": "文脈"
        },
        "ko": {
            "title": "회의 분석 보고서",
            "summary": "회의 요약",
            "topics": "주요 주제",
            "actions": "액션 아이템",
            "decisions": "중요한 결정",
            "task": "작업",
            "assignee": "담당자",
            "deadline": "마감일",
            "decision": "결정",
            "context": "맥락"
        },
        "zh": {
            "title": "会议分析报告",
            "summary": "会议摘要",
            "topics": "主要主题",
            "actions": "行动项",
            "decisions": "重要决定",
            "task": "任务",
            "assignee": "负责人",
            "deadline": "截止日期",
            "decision": "决定",
            "context": "背景"
        }
    }
    
    lang = labels.get(current_language, labels["vi"])
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading(lang["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Summary
    doc.add_heading(lang["summary"], 1)
    doc.add_paragraph(last_summary)

    # Topics
    if last_topics:
        doc.add_heading(lang["topics"], 1)
        for i, topic in enumerate(last_topics, 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{topic.get('topic', 'N/A')}").bold = True
            doc.add_paragraph(topic.get('description', ''), style='List Bullet 2')

    # Action Items
    if last_actions:
        doc.add_heading(lang["actions"], 1)
        for i, action in enumerate(last_actions, 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{action.get('task', 'N/A')}").bold = True
            doc.add_paragraph(f"{lang['assignee']}: {action.get('assignee', 'N/A')}", style='List Bullet 2')
            doc.add_paragraph(f"{lang['deadline']}: {action.get('deadline', 'N/A')}", style='List Bullet 2')

    # Decisions
    if last_decisions:
        doc.add_heading(lang["decisions"], 1)
        for i, decision in enumerate(last_decisions, 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{decision.get('decision', 'N/A')}").bold = True
            doc.add_paragraph(f"{lang['context']}: {decision.get('context', 'N/A')}", style='List Bullet 2')

    # Save to file
    filename = f"meeting_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = Path(filename)
    doc.save(str(filepath))

    return str(filepath)


# Simple and clean Gradio interface
with gr.Blocks(title="Meeting Transcript Chatbot") as demo:

    gr.Markdown("""
    # 💬 Meeting Transcript Chatbot
    ### Phân tích cuộc họp bằng AI
    """)

    # History section
    with gr.Accordion("📚 Lịch sử phân tích", open=False):
        with gr.Row():
            history_dropdown = gr.Dropdown(
                label="Chọn phân tích đã lưu",
                choices=[],
                interactive=True,
                scale=3
            )
            refresh_history_btn = gr.Button("🔄 Làm mới", scale=1)
            load_history_btn = gr.Button("📂 Tải lại", variant="primary", scale=1)

        history_info = gr.Markdown("_Chưa có lịch sử_")



    # Step 1: Upload
    gr.Markdown("## 📤 Bước 1: Upload File Transcript")
    with gr.Row():
        file_input = gr.File(
            label="Chọn file TXT hoặc DOCX",
            file_types=[".txt", ".docx"]
        )
        process_btn = gr.Button("🚀 Xử lý Transcript", variant="primary", size="lg")

    status_output = gr.Textbox(label="Trạng thái", interactive=False)

    gr.Markdown("---")

    # Step 2: View Results
    gr.Markdown("## 📊 Bước 2: Xem Kết quả Phân tích")

    with gr.Row():
        with gr.Column():
            gr.Markdown("### 📝 Tóm tắt Cuộc họp")
            summary_output = gr.Textbox(
                lines=5,
                interactive=False,
                placeholder="Tóm tắt sẽ hiển thị ở đây sau khi xử lý..."
            )

        with gr.Column():
            gr.Markdown("### ✅ Action Items")
            actions_output = gr.Markdown("_Chưa có dữ liệu_")

    with gr.Row():
        with gr.Column():
            gr.Markdown("### 🎯 Chủ đề Chính")
            topics_output = gr.Markdown("_Chưa có dữ liệu_")

        with gr.Column():
            gr.Markdown("### 🎯 Quyết định Quan trọng")
            decisions_output = gr.Markdown("_Chưa có dữ liệu_")

    gr.Markdown("---")

    # Step 3: Chat with AI (Function Calling)
    gr.Markdown("## 💬 Bước 3: Hỏi Đáp với AI (Function Calling)")

    with gr.Row():
        with gr.Column(scale=2):
            chatbot_display = gr.Chatbot(
                label="Cuộc trò chuyện",
                height=400,
                show_label=True
            )
            with gr.Row():
                chat_input = gr.Textbox(
                    label="Câu hỏi của bạn",
                    placeholder="Ví dụ: Alice có những task gì? / Tìm từ 'budget' trong transcript / Ai tham gia meeting?",
                    scale=4
                )
                chat_btn = gr.Button("📤 Gửi", variant="primary", scale=1)
            
            clear_chat_btn = gr.Button("🗑️ Xóa lịch sử chat", variant="secondary", size="sm")
        
        with gr.Column(scale=1):
            gr.Markdown("""
            ### 🔧 Functions Available
            
            AI có thể tự động gọi các functions:
            
            **1. extract_action_items**
            - Trích xuất tasks
            - Filter theo người
            
            **2. get_meeting_participants**
            - Danh sách người tham gia
            - Role của từng người
            
            **3. search_transcript**
            - Tìm keyword
            - Hiển thị context
            
            **4. extract_decisions**
            - Quyết định quan trọng
            - Bối cảnh quyết định
            
            ---
            
            **💡 Ví dụ câu hỏi:**
            - "Alice có task gì?"
            - "Tìm từ 'budget'"
            - "Ai tham gia meeting?"
            - "Quyết định gì được đưa ra?"
            """)
    
    gr.Markdown("---")
    
    # Step 4: Export Results
    gr.Markdown("## 💾 Bước 4: Xuất Kết quả")
    
    with gr.Row():
        export_txt_btn = gr.Button("📄 Xuất file TXT", variant="secondary", size="lg", scale=1)
        export_docx_btn = gr.Button("📝 Xuất file DOCX", variant="secondary", size="lg", scale=1)
    
    export_output = gr.File(label="File đã xuất")
    
    gr.Markdown("""
    ---
    **💡 Hướng dẫn:** Upload file → Xử lý → Xem kết quả → Chat với AI → Xuất file
    """)
    
    # Event handlers
    
    # History handlers
    refresh_history_btn.click(
        fn=refresh_history,
        outputs=[history_dropdown, history_info]
    )
    
    load_history_btn.click(
        fn=load_history,
        inputs=[history_dropdown],
        outputs=[status_output, summary_output, topics_output, actions_output, decisions_output]
    )
    
    # Auto-refresh history on page load
    demo.load(
        fn=refresh_history,
        outputs=[history_dropdown, history_info]
    )
    
    process_btn.click(
        fn=process_file,
        inputs=[file_input],
        outputs=[status_output, summary_output, topics_output, actions_output, decisions_output]
    ).then(
        fn=refresh_history,  # Refresh history after processing
        outputs=[history_dropdown, history_info]
    )
    
    # Chat handlers
    chat_btn.click(
        fn=chat_with_ai,
        inputs=[chat_input, chatbot_display],
        outputs=[chatbot_display]
    ).then(
        fn=lambda: "",  # Clear input after sending
        outputs=[chat_input]
    )
    
    chat_input.submit(  # Allow Enter key to send
        fn=chat_with_ai,
        inputs=[chat_input, chatbot_display],
        outputs=[chatbot_display]
    ).then(
        fn=lambda: "",
        outputs=[chat_input]
    )
    
    clear_chat_btn.click(
        fn=clear_chat,
        outputs=[chatbot_display]
    )
    
    export_txt_btn.click(
        fn=export_to_txt,
        outputs=[export_output]
    )
    
    export_docx_btn.click(
        fn=export_to_docx,
        outputs=[export_output]
    )


if __name__ == "__main__":
    demo.launch(
        server_name="0.0.0.0",
        server_port=7862,
        share=False
    )
